import os
import re
import sys
import math
import json
import time
import uuid
import pickle
import base64
import shutil
import subprocess
import tempfile
import hashlib
import hmac
import logging
import secrets
import ipaddress
import socket
import threading
import urllib.request
import urllib.parse
import urllib.error
import html.parser
from datetime import datetime, timezone
from pathlib import Path
from collections import Counter
from typing import Annotated, Any
from filelock import FileLock

try:
    import docx  # type: ignore
except ImportError:
    docx = None

import pymupdf as fitz  # PyMuPDF — imports the real module directly, avoiding
                        # the deprecated `import fitz` legacy-alias shim that
                        # prints the runtime warning. Every fitz.* call below
                        # is unaffected since this is just a local alias.
from fastapi import Body, Depends, FastAPI, File, Form, Query, Request, Response, UploadFile
from fastapi.encoders import jsonable_encoder
from fastapi.exceptions import RequestValidationError
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, ConfigDict
from starlette.datastructures import Headers
from starlette.middleware import Middleware
from starlette.middleware.sessions import SessionMiddleware
from starlette.types import ASGIApp, Receive, Scope, Send
from werkzeug.utils import secure_filename  # Starlette does not vendor this; werkzeug is an explicit dep

from trace_store import (
    TraceStore, redact, redact_deep, register_prompt, get_prompt,
    QA_PROMPT_VERSION, RERANK_PROMPT_VERSION, REWRITE_PROMPT_VERSION,
)

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

def _load_env():
    env_path = Path(__file__).parent / ".env"
    if env_path.exists():
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                key, value = line.split("=", 1)
                os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


_load_env()


class ColoredFormatter(logging.Formatter):
    """Custom logging formatter that adds ANSI color codes and file:lineno to terminal log output."""
    COLORS = {
        logging.DEBUG: "\033[36m",      # Cyan
        logging.INFO: "\033[32m",       # Green
        logging.WARNING: "\033[33m",    # Yellow
        logging.ERROR: "\033[31m",      # Red
        logging.CRITICAL: "\033[41;1m", # White on Red
    }
    RESET = "\033[0m"
    BOLD = "\033[1m"
    DIM = "\033[2m"

    def format(self, record):
        color = self.COLORS.get(record.levelno, self.RESET)
        time_str = f"{self.DIM}{self.formatTime(record, '%Y-%m-%d %H:%M:%S')}{self.RESET}"
        level_str = f"{color}{self.BOLD}{record.levelname:7s}{self.RESET}"
        loc_str = f"{self.DIM}[{record.name}:{record.filename}:{record.lineno}]{self.RESET}"
        msg = record.getMessage()
        return f"{time_str} {level_str} {loc_str} {color}{msg}{self.RESET}"


def setup_logger():
    level_name = os.environ.get("LOG_LEVEL", "INFO").upper()
    log_level = getattr(logging, level_name, logging.INFO)
    log = logging.getLogger("ask_my_docs")
    log.setLevel(log_level)
    if not log.handlers:
        handler = logging.StreamHandler(sys.stdout)
        handler.setLevel(log_level)
        handler.setFormatter(ColoredFormatter())
        log.addHandler(handler)
        log.propagate = False
    return log


logger = setup_logger()


class RetrievalBackendError(RuntimeError):
    """Raised when the vector database backend fails during retrieval."""
    pass


class RAGTracer:
    """Provides formatted step-by-step trace logging for document ingestion and retrieval pipelines."""

    @staticmethod
    def trace(pipeline: str, step: int, total_steps: int, name: str, details: dict):
        header = f"\033[1;35m[TRACE | {pipeline}] \033[1;36mStep {step}/{total_steps}: {name}\033[0m"
        lines = [header]
        for key, value in details.items():
            lines.append(f"  \033[33m├─ {key:<24}\033[0m: \033[1;32m{value}\033[0m")
        logger.info("\n" + "\n".join(lines))

# ─────────────────────────────────────────────────────────────────────────────
#  App Setup
# ─────────────────────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent

_env_secret = os.environ.get("SECRET_KEY")
# APP_ENV is the preferred name now that Flask is gone; ENV and the legacy
# FLASK_ENV are still honored so existing .env files keep working unchanged.
_env_mode = os.environ.get(
    "APP_ENV", os.environ.get("ENV", os.environ.get("FLASK_ENV", ""))
).lower()

if _env_secret:
    SECRET_KEY = _env_secret
elif _env_mode == "production":
    raise RuntimeError(
        "CRITICAL: SECRET_KEY environment variable is mandatory in production mode. "
        "Running with an ephemeral per-process key will break sessions across worker restarts."
    )
else:
    SECRET_KEY = secrets.token_hex(32)
    logger.warning(
        "SECRET_KEY not set — using a random per-process key for development. Sessions will "
        "NOT survive a server restart, and running multiple workers (e.g. gunicorn) "
        "will break sessions. Set SECRET_KEY explicitly for production deployments."
    )

# 50 MB per request. Flask enforced this itself via app.config["MAX_CONTENT_LENGTH"];
# Starlette/FastAPI impose NO body-size limit of their own, so MaxBodySizeMiddleware
# below re-implements the cap explicitly. Dropping it would turn every upload
# endpoint into an unbounded read.
MAX_CONTENT_LENGTH = 50 * 1024 * 1024
SESSION_COOKIE_MAX_AGE = 14 * 24 * 60 * 60  # 14 days, matching Starlette's default
# Off by default, exactly as Flask's SESSION_COOKIE_SECURE was. Turn it on
# ONLY once TLS actually terminates in front of the app: browsers silently
# drop a Secure cookie sent over plain HTTP, which breaks sessions rather
# than hardening them.
SESSION_COOKIE_SECURE = os.environ.get("SESSION_COOKIE_SECURE", "").lower() in ("1", "true", "yes")


class _BodyTooLarge(BaseException):
    """Raised from the wrapped `receive` once a request body exceeds the cap.

    Deliberately a BaseException, not an Exception: FastAPI wraps *any*
    Exception raised while parsing a request body into a generic
    400 "There was an error parsing the body", which would swallow the 413
    this is supposed to produce. MaxBodySizeMiddleware has a second
    safety net for that case anyway (see send_wrapper), but not tripping
    the wrapper in the first place keeps the common path honest.
    """


def _too_large_response() -> JSONResponse:
    return JSONResponse({"error": "File too large (max 50 MB)"}, status_code=413)


class MaxBodySizeMiddleware:
    """
    Enforces a hard cap on request body size — the FastAPI equivalent of
    Flask's MAX_CONTENT_LENGTH.

    Two layers, because neither alone is sufficient:
      1. A declared Content-Length larger than the cap is rejected before a
         single body byte is read. This catches every ordinary browser or
         HTTP-client upload.
      2. Bodies with no declared length (HTTP/1.1 chunked transfer encoding)
         are counted as they stream through and aborted the moment the
         running total crosses the cap — otherwise (1) is trivially bypassed
         by simply omitting the header.

    Pure ASGI rather than BaseHTTPMiddleware so the body can be inspected
    without buffering the whole request first, which would defeat the point.
    """

    def __init__(self, app: ASGIApp, max_body_size: int) -> None:
        self.app = app
        self.max_body_size = max_body_size

    def _declares_oversized_body(self, scope: Scope) -> bool:
        """True only when Content-Length is present, parseable, and over the
        cap. A missing or malformed header falls through to the byte counter
        rather than being trusted either way."""
        declared = Headers(scope=scope).get("content-length")
        if declared is None:
            return False
        try:
            return int(declared) > self.max_body_size
        except ValueError:
            return False

    async def __call__(self, scope: Scope, receive: Receive, send: Send) -> None:
        if scope["type"] != "http":
            await self.app(scope, receive, send)
            return

        if self._declares_oversized_body(scope):
            await _too_large_response()(scope, receive, send)
            return

        received = 0
        exceeded = False
        response_started = False
        replaced = False

        async def limited_receive() -> Any:
            nonlocal received, exceeded
            message = await receive()
            if message["type"] == "http.request":
                received += len(message.get("body", b""))
                if received > self.max_body_size:
                    exceeded = True
                    raise _BodyTooLarge()
            return message

        async def send_wrapper(message: Any) -> None:
            nonlocal response_started, replaced
            if replaced:
                return  # the rest of the superseded response is dropped
            if message["type"] == "http.response.start":
                if exceeded:
                    # Something downstream caught the aborted read and turned
                    # it into its own error response. Replace it: the caller
                    # must see 413, not a misleading 400/500.
                    replaced = True
                    await _too_large_response()(scope, receive, send)
                    return
                response_started = True
            await send(message)

        try:
            await self.app(scope, limited_receive, send_wrapper)
        except BaseException:
            if not exceeded:
                raise
            # The exception is a consequence of us aborting the read (possibly
            # re-wrapped by an intermediate task group), so it is ours to answer.
            if not response_started and not replaced:
                await _too_large_response()(scope, receive, send)


app = FastAPI(
    title="Ask My Docs",
    description=(
        "Universal multimodal RAG: upload documents, retrieve with hybrid "
        "BM25 + embedding search fused by RRF, and answer with grounded citations."
    ),
    version="5.0.0",
    middleware=[
        # Outermost first. The size cap runs before session decoding so an
        # oversized body is refused without any further work being done on it.
        Middleware(MaxBodySizeMiddleware, max_body_size=MAX_CONTENT_LENGTH),
        Middleware(
            SessionMiddleware,
            secret_key=SECRET_KEY,
            session_cookie="session",
            max_age=SESSION_COOKIE_MAX_AGE,
            same_site="lax",
            https_only=SESSION_COOKIE_SECURE,
        ),
    ],
)

FRONTEND_DIR = BASE_DIR / "frontend"
FRONTEND_DIST = FRONTEND_DIR / "dist"

if (FRONTEND_DIST / "assets").exists():
    app.mount("/assets", StaticFiles(directory=str(FRONTEND_DIST / "assets")), name="assets")


def ensure_frontend_built(force: bool = False) -> bool:
    """
    Ensures that the React 18 + TypeScript + Vite production build exists.
    If frontend/dist/index.html is missing (or force=True), executes:
      npm.cmd run build (on Windows) or npm run build (on Linux/macOS)
    from the frontend/ directory.

    Returns True if the build exists or succeeded.
    Raises SystemExit(1) if the build fails, npm is missing, or index.html is not created.
    """
    index_html = FRONTEND_DIST / "index.html"
    if index_html.exists() and not force:
        return True

    package_json = FRONTEND_DIR / "package.json"
    if not package_json.exists():
        print(f"❌ frontend/package.json not found at {FRONTEND_DIR}", file=sys.stderr)
        logger.error("frontend/package.json not found at %s", FRONTEND_DIR)
        raise SystemExit(1)

    npm_cmd = "npm.cmd" if sys.platform.startswith("win") else "npm"
    logger.info("📦 React frontend not built. Executing '%s run build' in %s...", npm_cmd, FRONTEND_DIR)
    print(f"\n📦 Building React 18 + Vite frontend ({npm_cmd} run build)...")

    try:
        res = subprocess.run(
            [npm_cmd, "run", "build"],
            cwd=str(FRONTEND_DIR),
            capture_output=True,
            text=True,
            check=False,
        )
        if res.returncode != 0:
            print(f"❌ Frontend build failed (exit code {res.returncode}):\n", file=sys.stderr)
            if res.stdout:
                print(res.stdout, file=sys.stderr)
            if res.stderr:
                print(res.stderr, file=sys.stderr)
            logger.error("Frontend build failed:\n%s\n%s", res.stdout, res.stderr)
            raise SystemExit(1)

        if not index_html.exists():
            print("❌ Frontend build finished with code 0 but frontend/dist/index.html was not generated.", file=sys.stderr)
            logger.error("Frontend build did not create %s", index_html)
            raise SystemExit(1)

        print("✅ Frontend build completed successfully!\n")
        logger.info("✅ Frontend build completed.")

        # Mount /assets if it wasn't mounted yet at startup
        assets_dir = FRONTEND_DIST / "assets"
        if assets_dir.exists():
            has_assets_route = any(getattr(route, "path", None) == "/assets" for route in app.routes)
            if not has_assets_route:
                app.mount("/assets", StaticFiles(directory=str(assets_dir)), name="assets")

        return True
    except FileNotFoundError:
        print(
            f"❌ '{npm_cmd}' was not found in PATH.\n"
            f"   Node.js and npm are required to build the frontend when frontend/dist is missing.\n"
            f"   Please install Node.js 18+ or run 'npm run build' inside frontend/.",
            file=sys.stderr,
        )
        logger.error("'%s' not found in PATH when building frontend.", npm_cmd)
        raise SystemExit(1)
    except SystemExit:
        raise
    except Exception as exc:
        print(f"❌ Unexpected error while building frontend: {exc}", file=sys.stderr)
        logger.error("Unexpected error during frontend build: %s", exc, exc_info=True)
        raise SystemExit(1)

# ─────────────────────────────────────────────────────────────────────────────
#  Trace logging (W5 Task Set C — durable, replayable /ask traces)
# ─────────────────────────────────────────────────────────────────────────────
# TRACE_LOG_PATH lets you point separate environments (or an eval run vs a
# real-traffic run) at separate trace files, without touching code.
TRACE_LOG_PATH = os.environ.get("TRACE_LOG_PATH", str(BASE_DIR / "traces" / "traces.jsonl"))
TRACES = TraceStore(TRACE_LOG_PATH)
UPLOAD_FOLDER = BASE_DIR / "uploads"
VECTOR_FOLDER = BASE_DIR / "vectorstore"
UPLOAD_FOLDER.mkdir(exist_ok=True)
VECTOR_FOLDER.mkdir(exist_ok=True)
ALLOWED_EXTENSIONS = {"pdf", "txt", "md"}

# Per-session stores
VECTOR_STORE: dict[str, "VectorStore"] = {}
SESSION_ACCESS: dict[str, float] = {}
HASH_STORE: dict[str, set[tuple[str, str]]] = {}
# Tracks which (content-hash, chunk_mode) key belongs to which doc_id, per
# session, so that removing a single document can also forget that key (see
# remove_doc()). Keying on (hash, chunk_mode) instead of hash alone means the
# exact same file can be indexed under several chunking strategies at once —
# needed to actually compare retrieval quality across chunk sizes, instead of
# only ever seeing a hypothetical chunk-count preview for strategies you
# never actually indexed. Uploading the same file + same chunk_mode twice is
# still rejected as a duplicate.
HASH_BY_DOC: dict[str, dict[str, tuple[str, str]]] = {}
SESSION_FILES: dict[str, dict[str, dict]] = {}
CHUNK_COUNTS: dict[str, dict[str, int]] = {}

# ─────────────────────────────────────────────────────────────────────────────
#  Durable Orphaned Document Registry (Operational Reconciliation)
# ─────────────────────────────────────────────────────────────────────────────
# When vector store rollback fails during upload or cancellation, the document
# is durably logged to orphans.jsonl so the failure is preserved across process
# restarts and observable across multiple worker processes.
ORPHAN_LOG_PATH = Path(os.environ.get("ORPHAN_LOG_PATH", str(UPLOAD_FOLDER / "orphans.jsonl")))
_ORPHAN_LOCK = threading.Lock()
ORPHANED_DOCS: dict[str, list[dict]] = {}  # In-memory index: sid -> list of orphan records
ADMIN_API_KEY = os.environ.get("ADMIN_API_KEY", "").strip()


def _is_admin_request(req: Request) -> bool:
    """Validate administrative authorization via ADMIN_API_KEY (Bearer token or X-Admin-Key)."""
    if not ADMIN_API_KEY:
        return False
    auth_header = req.headers.get("Authorization", "")
    token = ""
    if auth_header.startswith("Bearer "):
        token = auth_header[7:].strip()
    admin_header = req.headers.get("X-Admin-Key", "").strip()
    candidate = token or admin_header
    if not candidate:
        return False
    return hmac.compare_digest(candidate, ADMIN_API_KEY)


def _get_orphan_lock(path: Path | None = None) -> FileLock:
    """Acquire a cross-process file lock for the orphan log."""
    target = path or ORPHAN_LOG_PATH
    lock_path = str(target) + ".lock"
    return FileLock(lock_path, timeout=10)


def _record_orphaned_doc(sid: str, doc_id: str, filename: str, error: str, stored_path: Path | str | None = None) -> dict:
    """Durably record an orphaned document in orphans.jsonl and in-memory registry.
    If durable file persistence fails, emit CRITICAL operational failure and flag
    reconciliation_persistence_failed=True."""
    record = {
        "session_id": sid,
        "doc_id": doc_id,
        "filename": filename,
        "error": str(error),
        "stored_path": str(stored_path) if stored_path else None,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "status": "orphaned",
    }
    with _ORPHAN_LOCK:
        try:
            ORPHAN_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
            with _get_orphan_lock():
                with ORPHAN_LOG_PATH.open("a", encoding="utf-8") as f:
                    f.write(json.dumps(record, ensure_ascii=False) + "\n")
                    f.flush()
                    if hasattr(os, "fsync"):
                        try:
                            os.fsync(f.fileno())
                        except OSError as fsync_err:
                            raise OSError(f"Orphan log fsync failed: {fsync_err}") from fsync_err
        except Exception as exc:
            record["reconciliation_persistence_failed"] = True
            logger.critical(
                "❌ CRITICAL OPERATIONAL FAILURE: Failed to write orphan reconciliation record to %s "
                "for doc %s (session %s): %s. Durable reconciliation tracking failed!",
                ORPHAN_LOG_PATH, doc_id, sid, exc, exc_info=True,
            )

        # Update in-memory registry
        ORPHANED_DOCS.setdefault(sid, []).append(record)
    return record


def _resolve_orphaned_doc(sid: str, doc_id: str) -> bool:
    """Mark an orphan record as resolved when successful cleanup occurs,
    appending the resolution event durably so restarts do not resurrect it.
    Durable resolution is the authoritative state transition: the orphan
    is only removed from in-memory cache after durable persistence succeeds."""
    with _ORPHAN_LOCK:
        in_memory_match = any(o.get("doc_id") == doc_id for o in ORPHANED_DOCS.get(sid, []))
        if not in_memory_match:
            durable_records = _read_durable_orphans()
            if not any(o.get("doc_id") == doc_id for o in durable_records.get(sid, [])):
                return False

        record = {
            "session_id": sid,
            "doc_id": doc_id,
            "status": "resolved",
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

        # Step 1: Durably append resolution record first
        try:
            ORPHAN_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
            with _get_orphan_lock():
                with ORPHAN_LOG_PATH.open("a", encoding="utf-8") as f:
                    f.write(json.dumps(record, ensure_ascii=False) + "\n")
                    f.flush()
                    if hasattr(os, "fsync"):
                        try:
                            os.fsync(f.fileno())
                        except OSError as fsync_err:
                            raise OSError(f"Orphan log fsync failed during resolution: {fsync_err}") from fsync_err
        except Exception as exc:
            logger.critical(
                "❌ CRITICAL OPERATIONAL FAILURE: Failed to write orphan resolution to %s for doc %s (session %s): %s. "
                "Orphan remains active in memory so reconciliation is not falsely dropped!",
                ORPHAN_LOG_PATH, doc_id, sid, exc, exc_info=True,
            )
            return False

        # Step 2: ONLY mutate in-memory cache after durable persistence succeeds
        if sid in ORPHANED_DOCS:
            ORPHANED_DOCS[sid] = [o for o in ORPHANED_DOCS[sid] if o.get("doc_id") != doc_id]

        return True


def _read_durable_orphans(path: Path | None = None) -> dict[str, list[dict]]:
    """Read and validate all durable orphan log records from disk,
    replaying additions and resolutions to produce the current unresolved orphan state.
    This guarantees cross-worker visibility by reading directly from shared durable storage."""
    log_path = path or ORPHAN_LOG_PATH
    if not log_path.exists():
        return {}

    orphans_by_sid: dict[str, list[dict]] = {}
    try:
        with _get_orphan_lock(log_path):
            with log_path.open("r", encoding="utf-8") as f:
                for line_no, line in enumerate(f, 1):
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        record = json.loads(line)
                    except json.JSONDecodeError as err:
                        logger.warning("Corrupted orphan log line %d in %s: %s", line_no, log_path, err)
                        continue

                    if not isinstance(record, dict):
                        logger.warning("Malformed orphan record (not an object) on line %d in %s", line_no, log_path)
                        continue

                    sid = record.get("session_id")
                    doc_id = record.get("doc_id")
                    status = record.get("status")

                    # Strict schema validation: non-empty session_id, doc_id, and recognized status
                    if not isinstance(sid, str) or not sid.strip() or not isinstance(doc_id, str) or not doc_id.strip():
                        logger.warning(
                            "Quarantining malformed orphan record on line %d in %s "
                            "(missing or invalid session_id/doc_id): %s",
                            line_no, log_path, record,
                        )
                        continue

                    if status == "resolved":
                        if sid in orphans_by_sid:
                            orphans_by_sid[sid] = [o for o in orphans_by_sid[sid] if o.get("doc_id") != doc_id]
                    elif status == "orphaned":
                        orphans_by_sid.setdefault(sid, []).append(record)
                    else:
                        logger.warning(
                            "Quarantining orphan record on line %d in %s with unknown status '%s': %s",
                            line_no, log_path, status, record,
                        )
                        continue
    except Exception as exc:
        logger.error("Failed to read durable orphan records from %s: %s", log_path, exc, exc_info=True)

    return orphans_by_sid


def _load_orphaned_docs() -> None:
    """Load durable orphan records from disk into the in-memory cache on startup."""
    with _ORPHAN_LOCK:
        ORPHANED_DOCS.clear()
        durable = _read_durable_orphans()
        for sid, records in durable.items():
            ORPHANED_DOCS[sid] = list(records)


# Maps a client-generated upload_id -> the time it was cancelled. /upload
# checks this (see below) so a mid-flight cancel can (a) stop processing
# any remaining files in a multi-file batch, and (b) roll back whatever
# was already added to the store during THIS call before the response
# goes out. Attempt to stop/rollback this upload's work. Backend deletion can fail,
# so incomplete cleanup is explicitly reported and reconciliation metadata
# is retained when possible.
CANCELLED_UPLOADS: dict[str, float] = {}
CANCELLED_UPLOAD_TTL = 10 * 60  # forget stale cancel signals after 10 minutes


def _sweep_cancelled_uploads() -> None:
    now = time.time()
    for uid in [u for u, t in CANCELLED_UPLOADS.items() if now - t > CANCELLED_UPLOAD_TTL]:
        CANCELLED_UPLOADS.pop(uid, None)

SESSION_TTL = 60 * 60  # 1 hour
MAX_SESSIONS = 20

# ─────────────────────────────────────────────────────────────────────────────
#  Google Gemini config (embeddings + LLM)
# ─────────────────────────────────────────────────────────────────────────────

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta"
EMBED_MODEL = os.environ.get("EMBED_MODEL", "gemini-embedding-001")
# Vision OCR (extract_image_pages) still calls Gemini directly for
# multimodal image understanding — xAI's chat-completions endpoint isn't
# used for that path, so it needs its OWN model constant, separate from
# XAI_MODEL/LLM_MODEL below (which now default to a text-only Grok model
# and would 400 if pointed at Gemini's generateContent endpoint).
GEMINI_VISION_MODEL = os.environ.get("GEMINI_VISION_MODEL", "gemini-3.7-flash")
# NOTE: Gemini is now used for embeddings (EMBED_MODEL) and image Vision
# OCR (GEMINI_VISION_MODEL) only. Every text chat/generation call (answer
# synthesis, reranking, query rewriting) goes to xAI's Grok models below —
# xAI does not currently offer a public embeddings API, so embeddings have
# nowhere else to go without also swapping the vector store's dimensionality.

# ── Embeddings backend selector ─────────────────────────────────────────
# "gemini" = Gemini's batchEmbedContents API — free tier has a real
# requests-per-minute quota, which is exactly what "HTTP 429 Too Many
# Requests" during a big upload means: not a bug, just the free-tier
# throughput ceiling for a 300+ chunk document going out in rapid batches.
# "ollama" (default here) = a locally-running Ollama server — no API
# key, no rate limit, no per-token cost, since it's your own machine
# doing the inference. Requires `ollama serve` running and the model
# pulled first:
#   ollama pull nomic-embed-text
# Switching AFTER you've already indexed documents under the other
# backend requires clearing and re-uploading — the two backends produce
# different-dimension vectors (Gemini: 3072-dim, nomic-embed-text:
# 768-dim) that live in incompatible vector spaces; Qdrant infers its
# collection's dimension from whichever vectors arrive first.
EMBED_BACKEND = os.environ.get("EMBED_BACKEND", "ollama")  # "gemini" | "ollama"
OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://localhost:11434")
OLLAMA_EMBED_MODEL = os.environ.get("OLLAMA_EMBED_MODEL", "nomic-embed-text")

# ── Vision OCR (image uploads) backend selector ──────────────────────────
# "gemini" uses GEMINI_API_KEY + GEMINI_VISION_MODEL above. "ollama"
# (default here) uses a local vision-capable Ollama model instead — same
# no-key, no-rate-limit, runs-on-your-machine tradeoff as embeddings.
#   ollama pull llava             (7B, solid general-purpose default)
#   ollama pull moondream         (1.8B, much lighter/faster, less accurate)
#   ollama pull llama3.2-vision   (stronger OCR/chart reading, needs more RAM)
# Local vision models are noticeably weaker at dense-text OCR than
# Gemini's — expect more missed/garbled text on small print or busy
# scans. Worth checking output quality on a real image before relying on it.
VISION_BACKEND = os.environ.get("VISION_BACKEND", "ollama")  # "gemini" | "ollama"
OLLAMA_VISION_MODEL = os.environ.get("OLLAMA_VISION_MODEL", "llava")

# ── xAI (Grok) config — used for chat/generation when CHAT_BACKEND=xai ──
XAI_API_KEY = os.environ.get("XAI_API_KEY", "")
XAI_URL = os.environ.get("XAI_URL", "https://api.x.ai/v1")
XAI_MODEL = os.environ.get("XAI_MODEL", "grok-4.3")

# ── Chat / generation backend selector ───────────────────────────────────
# "xai" = xAI's Grok models via the cloud API (needs XAI_API_KEY, costs
# per token, but noticeably stronger at grounded citation-following and
# staying strictly within the retrieved excerpts). "ollama" = a fully
# local chat model instead — no API key, no cost, nothing leaves your
# machine, but a locally-runnable model is meaningfully weaker at
# instruction-following (expect looser citation discipline and more
# hallucination) and can be much slower without a real GPU.
#   ollama pull llama3.1     (8B, solid general default)
#   ollama pull qwen2.5      (good instruction-following for its size)
#   ollama pull mistral      (7B, fast, lighter weight)
CHAT_BACKEND = os.environ.get("CHAT_BACKEND", "ollama")  # "xai" | "ollama"
OLLAMA_CHAT_MODEL = os.environ.get("OLLAMA_CHAT_MODEL", "llama3.1")

# LLM_MODEL is kept as the generic "which chat model is answering right
# now" name used everywhere else in this file (trace logs, RAGTracer,
# /status) — it now follows whichever CHAT_BACKEND is active.
LLM_MODEL = OLLAMA_CHAT_MODEL if CHAT_BACKEND == "ollama" else XAI_MODEL
EMBED_MIN_SCORE = float(os.environ.get("EMBED_MIN_SCORE", "0.55"))  # recalibrated
# for reciprocal_rank_fusion()'s normalized scale (see RRF_K comment) —
# NOT comparable to the old weighted-average blend's raw-cosine scale.
# Tested directly: a genuine top-ranked match scores ~0.9-1.0 here; a
# plausible false positive (moderate, non-top rank in both methods)
# scores ~0.39-0.42 — 0.55 sits cleanly between them.
EMBED_BATCH = 32

# Retrieval strategy for the embeddings path. "hybrid" combines cosine
# similarity with TF-IDF so exact tokens (error codes, IDs, acronyms) aren't
# lost to semantic blurring. "embed" is the old behavior, kept only so you
# can flip back to it and measure the before/after with the same code.
RETRIEVAL_MODE = os.environ.get("RETRIEVAL_MODE", "hybrid")  # "hybrid" | "embed"
HYBRID_ALPHA = float(os.environ.get("HYBRID_ALPHA", "0.6"))  # weight on embedding score

# How many chunks retrieval returns, before token-budget trimming below.
TOP_K = int(os.environ.get("TOP_K", "8"))  # was 5 — a genuinely relevant
# section that's split across chunks (see STRUCT_MAX_WORDS) can get
# crowded out of a small top-k window entirely by other, unrelated but
# similarly-scored sections. A wider window gives the LLM more raw
# material to synthesize a complete answer from.

# Hard cap on how many tokens' worth of retrieved chunk TEXT get sent to
# the LLM as context, regardless of how many chunks TOP_K allows through.
#
# Why this is needed even though TOP_K already exists: TOP_K caps the
# COUNT of chunks, not their combined SIZE. Structured chunking has no
# hard upper bound on a single chunk (see structured_chunk() /
# split_by_sentences() — an unstructured document with no real sentence
# punctuation can produce one abnormally large chunk). Even 5 large
# chunks — or just one — can already blow past a reasonable context
# budget, drive up latency/cost, and in the worst case exceed the model's
# actual context window. This trims the already-retrieved, already-ranked
# list down to what actually fits, on top of (not instead of) TOP_K.
MAX_CONTEXT_TOKENS = int(os.environ.get("MAX_CONTEXT_TOKENS", "6000"))  # was
# 3000 — raised alongside TOP_K (5->8); otherwise the wider candidate
# pool just gets trimmed straight back down by this budget, defeating
# the point of considering more candidates in the first place.

# Which VectorStore implementation backs retrieval. "memory" (default) is
# the original brute-force in-process store — zero setup, pickled to disk.
# "qdrant" swaps in a real vector database with an actual ANN index (HNSW)
# behind it. The SAME code path handles both self-hosted Qdrant (via
# docker-compose.yml, QDRANT_API_KEY left unset) and Qdrant Cloud (a
# cluster URL + API key from cloud.qdrant.io) — only QDRANT_URL and
# QDRANT_API_KEY differ between the two; see qdrant_store.py.
VECTOR_BACKEND = os.environ.get("VECTOR_BACKEND", "memory")
QDRANT_URL = os.environ.get("QDRANT_URL", "http://localhost:6333")
QDRANT_API_KEY = os.environ.get("QDRANT_API_KEY") or None  # None for self-hosted, unauthenticated
QDRANT_TIMEOUT = float(os.environ.get("QDRANT_TIMEOUT", "10"))
# ANN candidates pulled from Qdrant before TF-IDF re-ranks for hybrid
# search — see qdrant_store.py's module docstring for why this differs
# from the in-memory backend's "score every chunk" approach.
QDRANT_CANDIDATE_POOL = int(os.environ.get("QDRANT_CANDIDATE_POOL", "30"))


def estimate_tokens(text: str) -> int:
    """
    Rough token-count estimate. There's no real tokenizer available
    offline for arbitrary models, so this uses the standard ~4-characters-
    per-token rule of thumb for English text (the same approximation most
    LLM providers themselves suggest when a real tokenizer isn't handy).
    Not exact — good enough for a budget guard, not for billing.
    """
    return max(1, len(text) // 4)


def fit_to_token_budget(results: list[dict], max_tokens: int) -> list[dict]:
    """
    Greedily keep results — assumed already sorted best-first by score —
    until the cumulative estimated token count would exceed max_tokens.

    Always keeps at least the single best-scoring result, even if it
    alone exceeds the budget: an answer grounded in one long chunk is
    still better than refusing to answer at all. Everything after that
    is dropped once the running total would tip over the limit, so a
    handful of oversized chunks can't silently balloon the prompt sent to
    the LLM regardless of how large TOP_K is set to.
    """
    kept: list[dict] = []
    used = 0
    for r in results:
        t = estimate_tokens(r.get("text", ""))
        if kept and used + t > max_tokens:
            break
        kept.append(r)
        used += t
    return kept

_simple_tokenizer = re.compile(r"[\w]+")


def _xai_chat_call(system: str, user: str,
                   temperature: float = 0,
                   max_tokens: int | None = None,
                   model: str | None = None) -> str:
    """Call xAI's OpenAI-compatible /v1/chat/completions endpoint, with
    automatic exponential backoff retry on HTTP 429/500/503.
    """
    url = f"{XAI_URL}/chat/completions"
    target_model = model or XAI_MODEL
    payload: dict = {
        "model": target_model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": temperature,
    }
    if max_tokens is not None:
        payload["max_tokens"] = max_tokens

    max_retries = 5
    base_delay = 2.0
    for attempt in range(1, max_retries + 1):
        try:
            t0 = time.time()
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {XAI_API_KEY}",
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=90) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            elapsed = time.time() - t0
            choices = data.get("choices", [])
            if not choices:
                logger.warning("⚠️ xAI response has no choices (raw: %s)", data)
                return ""
            result_text = (choices[0].get("message", {}).get("content") or "").strip()
            logger.info("🤖 xAI (Grok) LLM call succeeded in %.2fs (model: %s)", elapsed, target_model)
            return result_text
        except urllib.error.HTTPError as exc:
            if exc.code in (429, 500, 503) and attempt < max_retries:
                delay = base_delay * (2 ** (attempt - 1))
                logger.warning("⏳ xAI LLM rate-limited (HTTP %d). Retrying in %.1fs (attempt %d/%d)...",
                               exc.code, delay, attempt, max_retries)
                time.sleep(delay)
            else:
                logger.error("❌ xAI LLM call failed with HTTP %d: %s", exc.code, exc)
                raise
        except (urllib.error.URLError, socket.timeout, TimeoutError, OSError) as exc:
            if attempt < max_retries:
                delay = base_delay * (2 ** (attempt - 1))
                logger.warning("⏳ xAI LLM network timeout (%s). Retrying in %.1fs (attempt %d/%d)...",
                               exc, delay, attempt, max_retries)
                time.sleep(delay)
            else:
                logger.error("❌ xAI LLM network call timed out after %d attempts: %s", max_retries, exc)
                raise
        except Exception as exc:
            logger.error("❌ xAI LLM call exception: %s", exc, exc_info=True)
            raise


def _ollama_chat_call(system: str, user: str,
                      temperature: float = 0,
                      max_tokens: int | None = None,
                      model: str | None = None) -> str:
    """Call a locally-running Ollama chat model via its native /api/chat endpoint."""
    url = f"{OLLAMA_URL}/api/chat"
    target_model = model or OLLAMA_CHAT_MODEL
    options: dict = {"temperature": temperature}
    if max_tokens is not None:
        options["num_predict"] = max_tokens
    payload = {
        "model": target_model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "stream": False,
        "options": options,
    }
    try:
        t0 = time.time()
        req = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        elapsed = time.time() - t0
        result_text = (data.get("message", {}).get("content") or "").strip()
        logger.info("🤖 Ollama (%s) LLM call succeeded in %.2fs", target_model, elapsed)
        return result_text
    except urllib.error.URLError as exc:
        logger.error(
            "❌ Could not reach Ollama at %s — is `ollama serve` running and "
            "`%s` pulled? (%s)", OLLAMA_URL, target_model, exc,
        )
        raise
    except Exception as exc:
        logger.error("❌ Ollama LLM call exception: %s", exc, exc_info=True)
        raise


def _chat_call(system: str, user: str,
               temperature: float = 0,
               max_tokens: int | None = None,
               model: str | None = None) -> str:
    """Single entry point for every chat/generation call in this file —
    routes to xAI or a local Ollama model per CHAT_BACKEND."""
    if CHAT_BACKEND == "ollama":
        return _ollama_chat_call(system, user, temperature=temperature, max_tokens=max_tokens, model=model)
    return _xai_chat_call(system, user, temperature=temperature, max_tokens=max_tokens, model=model)


def _chat_configured() -> bool:
    """True if the selected chat backend is actually usable right now:
    CHAT_BACKEND=ollama needs no key (a local server is assumed reachable;
    a connection failure surfaces at call time rather than being
    pre-checked here); CHAT_BACKEND=xai needs XAI_API_KEY."""
    if CHAT_BACKEND == "ollama":
        return True
    return bool(XAI_API_KEY)


def _gemini_embed_batch(texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts via Gemini batchEmbedContents API with automatic backoff on HTTP 429/503."""
    url = f"{GEMINI_URL}/models/{EMBED_MODEL}:batchEmbedContents?key={GEMINI_API_KEY}"
    model_name = EMBED_MODEL if EMBED_MODEL.startswith("models/") else f"models/{EMBED_MODEL}"
    payload = {
        "requests": [
            {
                "model": model_name,
                "content": {"parts": [{"text": t}]},
            }
            for t in texts
        ]
    }

    max_retries = 5
    base_delay = 2.0
    for attempt in range(1, max_retries + 1):
        try:
            t0 = time.time()
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=90) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            elapsed = time.time() - t0
            vectors = [e["values"] for e in data["embeddings"]]
            logger.info("🧠 Gemini Embedding batch of %d chunks succeeded in %.2fs (vector dim: %d)",
                        len(texts), elapsed, len(vectors[0]) if vectors else 0)
            return vectors
        except urllib.error.HTTPError as exc:
            if exc.code in (429, 500, 503) and attempt < max_retries:
                delay = base_delay * (2 ** (attempt - 1))
                logger.warning("⏳ Gemini Embeddings rate-limited (HTTP %d). Retrying in %.1fs (attempt %d/%d)...",
                               exc.code, delay, attempt, max_retries)
                time.sleep(delay)
            else:
                logger.error("❌ Gemini Embeddings batch failed with HTTP %d: %s", exc.code, exc)
                raise
        except Exception as exc:
            logger.error("❌ Gemini Embeddings batch exception: %s", exc, exc_info=True)
            raise


def _ollama_embed_batch(texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts via a locally-running Ollama server's
    native /api/embed endpoint. No API key, no rate limit, no per-token
    cost — it's your own machine doing the inference. Requires
    `ollama serve` running and the embedding model already pulled
    (`ollama pull nomic-embed-text`)."""
    url = f"{OLLAMA_URL}/api/embed"
    payload = {"model": OLLAMA_EMBED_MODEL, "input": texts}
    try:
        t0 = time.time()
        req = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        elapsed = time.time() - t0
        vectors = data.get("embeddings", [])
        logger.info("🧠 Ollama embedding batch of %d chunks succeeded in %.2fs (vector dim: %d, model: %s)",
                    len(texts), elapsed, len(vectors[0]) if vectors else 0, OLLAMA_EMBED_MODEL)
        return vectors
    except urllib.error.URLError as exc:
        logger.error(
            "❌ Could not reach Ollama at %s — is `ollama serve` running and "
            "`%s` pulled? (%s)", OLLAMA_URL, OLLAMA_EMBED_MODEL, exc,
        )
        raise
    except Exception as exc:
        logger.error("❌ Ollama embedding batch exception: %s", exc, exc_info=True)
        raise


def embed_texts(texts: list[str]) -> list[list[float]]:
    """Batch-embed texts via whichever backend EMBED_BACKEND selects
    (Gemini API or a local Ollama server)."""
    out: list[list[float]] = []
    total_batches = math.ceil(len(texts) / EMBED_BATCH)
    backend_label = "Ollama (local)" if EMBED_BACKEND == "ollama" else "Gemini"
    logger.info("🧠 Starting batch embedding for %d total chunks (%d batch(es) of max %d) via %s",
                len(texts), total_batches, EMBED_BATCH, backend_label)
    for idx, i in enumerate(range(0, len(texts), EMBED_BATCH), start=1):
        batch = texts[i:i + EMBED_BATCH]
        logger.info("🧠 Processing embedding batch %d/%d (%d chunks)...", idx, total_batches, len(batch))
        vectors = (_ollama_embed_batch(batch) if EMBED_BACKEND == "ollama"
                   else _gemini_embed_batch(batch))
        out.extend(vectors)
    return out


def embed_text(text: str) -> list[float]:
    return embed_texts([text])[0]


def _embeddings_configured() -> bool:
    """True if the selected embeddings backend is actually usable right
    now: EMBED_BACKEND=ollama needs no key at all (a local server is
    assumed reachable — a connection failure surfaces at embed time
    instead of being pre-checked here, to avoid an extra network round
    trip on every call site); EMBED_BACKEND=gemini needs GEMINI_API_KEY."""
    if EMBED_BACKEND == "ollama":
        return True
    return bool(GEMINI_API_KEY)


# ─────────────────────────────────────────────────────────────────────────────
#  VectorStore (in-memory + disk-persisted)
# ─────────────────────────────────────────────────────────────────────────────

def _dot(a: list[float], b: list[float]) -> float:
    return sum(x * y for x, y in zip(a, b))


def _norm(a: list[float]) -> float:
    return math.sqrt(sum(x * x for x in a))


def cosine(a: list[float], b: list[float]) -> float:
    denom = _norm(a) * _norm(b)
    return 0.0 if denom == 0 else _dot(a, b) / denom


class VectorStore:
    """A tiny persistent vector index: chunk metadata + embedding vectors."""

    def __init__(self, sid: str):
        self.sid = sid
        self.chunks: list[dict] = []
        self.vectors: list[list[float]] = []
        self.path = VECTOR_FOLDER / f"{sid}.pkl"
        # Cached TF-IDF index (token frequencies + IDF weights) for this
        # store's current chunk set. None means "needs rebuilding" — see
        # get_tfidf_index(). Invalidated by every mutation below so it's
        # never stale, but never rebuilt more often than necessary either.
        self._tfidf_index_cache: dict | None = None
        self.last_backend_error: str | None = None  # always None for this backend; see QdrantVectorStore

    def load(self) -> None:
        if self.path.exists():
            try:
                data = pickle.loads(self.path.read_bytes())
                self.chunks = data["chunks"]
                self.vectors = data["vectors"]
                self._tfidf_index_cache = None
            except Exception:
                logger.warning("Corrupted vector store at %s — starting fresh", self.path, exc_info=True)
                self.chunks, self.vectors = [], []

    def save(self) -> None:
        self.path.write_bytes(pickle.dumps({"chunks": self.chunks, "vectors": self.vectors}))

    def add(self, chunks: list[dict], vectors: list[list[float]]) -> None:
        self.chunks.extend(chunks)
        self.vectors.extend(vectors)
        self._tfidf_index_cache = None
        self.save()

    def get_tfidf_index(self) -> dict:
        """
        Return this store's TF-IDF index, building it once and caching it
        until the next mutation (add/remove_doc/clear/load all invalidate
        the cache).

        Why: without this, build_index() re-tokenizes and recomputes IDF
        over the ENTIRE chunk corpus on every single question that uses
        the TF-IDF path — repeated, wasted work that scales with corpus
        size and gets worse the more documents are indexed. The index
        only actually changes when chunks change, so it's safe to reuse
        between questions.
        """
        if self._tfidf_index_cache is None:
            self._tfidf_index_cache = build_index(self.chunks)
        return self._tfidf_index_cache

    def query(self, vector: list[float], top_k: int = 5,
              min_score: float = 0.0) -> list[dict]:
        scored = [(i, cosine(vector, v)) for i, v in enumerate(self.vectors)]
        scored = [(i, s) for i, s in scored if s >= min_score]
        scored.sort(key=lambda x: -x[1])
        return [{**self.chunks[i], "score": s} for i, s in scored[:top_k]]

    def query_scores(self, vector: list[float]) -> list[float]:
        """Cosine similarity of `vector` against every chunk, unfiltered and
        in chunk order. Used by hybrid_search() to line embedding scores up
        against TF-IDF scores index-for-index before combining them."""
        return [cosine(vector, v) for v in self.vectors]

    def clear(self) -> None:
        self.chunks, self.vectors = [], []
        self._tfidf_index_cache = None
        self.path.unlink(missing_ok=True)

    def remove_doc(self, doc_id: str) -> int:
        """Remove all chunks + vectors for a doc. Returns how many were removed."""
        before = len(self.chunks)
        kept = [(c, v) for c, v in zip(self.chunks, self.vectors)
                if c["doc_id"] != doc_id]
        self.chunks = [c for c, _ in kept]
        self.vectors = [v for _, v in kept]
        removed = before - len(self.chunks)
        if removed:
            self._tfidf_index_cache = None
            self.save()
        return removed

    def filtered_by_method(self, method: str) -> "VectorStore":
        """
        Return an ephemeral, non-persisted view containing only chunks
        indexed under the given chunking strategy ("structured", "128",
        "256", "512").

        This is what makes chunk-size comparison real rather than
        hypothetical: once the same document has been uploaded under two
        or more strategies (see the (hash, chunk_mode) dedupe key in
        /upload), /ask can restrict retrieval to just one strategy at a
        time, so you can directly see whether a given question is found
        under 128-word chunks but missed under 512-word chunks, etc.

        Never call .save()/.add() on the returned view — it shares no
        state with the real store and isn't written to disk.
        """
        view = VectorStore(f"{self.sid}__view")
        # In TF-IDF-only mode (no embeddings configured) self.vectors is []
        # while self.chunks isn't — zip(chunks, vectors) would silently
        # truncate to the shorter list and drop every chunk. Filter by
        # index instead so this works whether or not vectors are present.
        has_vectors = len(self.vectors) == len(self.chunks)
        matched = [i for i, c in enumerate(self.chunks) if c.get("method") == method]
        view.chunks = [self.chunks[i] for i in matched]
        view.vectors = [self.vectors[i] for i in matched] if has_vectors else []
        return view


def _manifest_path(sid: str) -> Path:
    return UPLOAD_FOLDER / f"{sid}.manifest.json"


_MANIFEST_MTIMES: dict[str, float] = {}


def _save_session_manifest(sid: str) -> None:
    """Persist the doc_id -> file mapping so it survives a server restart and syncs across workers."""
    files = SESSION_FILES.get(sid)
    mpath = _manifest_path(sid)
    if files:
        tmp_path = mpath.with_suffix(".tmp")
        tmp_path.write_text(json.dumps({
            doc_id: {"path": str(info["path"]), "name": info["name"]}
            for doc_id, info in files.items()
        }), encoding="utf-8")
        tmp_path.replace(mpath)
        try:
            _MANIFEST_MTIMES[sid] = mpath.stat().st_mtime
        except OSError:
            pass
    else:
        mpath.unlink(missing_ok=True)
        _MANIFEST_MTIMES.pop(sid, None)


def _load_session_manifest(sid: str) -> None:
    """Restore the file mapping for a session from disk."""
    manifest = _manifest_path(sid)
    if not manifest.exists():
        return
    try:
        data = json.loads(manifest.read_text(encoding="utf-8"))
    except Exception:
        logger.warning("Corrupted session manifest at %s — ignoring", manifest, exc_info=True)
        return
    files: dict[str, dict] = {}
    for doc_id, meta in data.items():
        p = Path(meta.get("path", ""))
        if p.exists():
            files[doc_id] = {"path": p, "name": meta.get("name", p.name)}
    if files or sid in SESSION_FILES:
        SESSION_FILES[sid] = files


def _cleanup_session_files(sid: str) -> None:
    """Delete any uploaded files retained for a session."""
    files = SESSION_FILES.pop(sid, None)
    if files:
        for info in files.values():
            info["path"].unlink(missing_ok=True)
    _manifest_path(sid).unlink(missing_ok=True)


def _sweep_orphan_uploads() -> None:
    """Remove upload files not referenced by any session manifest."""
    referenced: set[str] = set()
    for manifest in UPLOAD_FOLDER.glob("*.manifest.json"):
        try:
            data = json.loads(manifest.read_text(encoding="utf-8"))
        except Exception:
            logger.warning("Corrupted manifest during startup sweep: %s — skipping", manifest, exc_info=True)
            continue
        for meta in data.values():
            referenced.add(str(Path(meta.get("path", "")).resolve()))
    for f in UPLOAD_FOLDER.iterdir():
        if f.name.endswith(".manifest.json") or f.name.endswith(".jsonl") or f.name.endswith(".lock"):
            continue
        if str(f.resolve()) not in referenced:
            f.unlink(missing_ok=True)


_sweep_orphan_uploads()
_load_orphaned_docs()


def _make_store(sid: str):
    """Construct + load the right VectorStore implementation for `sid`,
    based on VECTOR_BACKEND. Both implementations expose the same
    interface, so nothing downstream needs to know or care which one it
    got. The qdrant_store import is lazy — only needed, and only
    attempted, when VECTOR_BACKEND=qdrant is actually set."""
    if VECTOR_BACKEND == "qdrant":
        from qdrant_store import QdrantVectorStore
        store = QdrantVectorStore(sid)
    else:
        store = VectorStore(sid)
    store.load()
    return store


def _evict_session_store(sid: str) -> None:
    """Release backend-specific storage for an evicted/cleared session.
    Safe to call even if that session never used the currently-active
    backend — the pickle file simply won't exist, or the Qdrant call
    will find no matching collection."""
    (VECTOR_FOLDER / f"{sid}.pkl").unlink(missing_ok=True)
    if VECTOR_BACKEND == "qdrant":
        try:
            from qdrant_store import QdrantVectorStore
            QdrantVectorStore(sid).clear()
        except Exception:
            logger.warning("Failed to release Qdrant collection for evicted session %s", sid, exc_info=True)


def _get_store(sid: str) -> VectorStore:
    """Return this session's vector store, evicting stale/overflowing sessions."""
    now = time.time()
    for s in list(SESSION_ACCESS):
        if now - SESSION_ACCESS[s] > SESSION_TTL:
            VECTOR_STORE.pop(s, None)
            SESSION_ACCESS.pop(s, None)
            HASH_STORE.pop(s, None)
            HASH_BY_DOC.pop(s, None)
            _evict_session_store(s)
            _cleanup_session_files(s)
    if len(SESSION_ACCESS) >= MAX_SESSIONS and sid not in SESSION_ACCESS:
        oldest = min(SESSION_ACCESS, key=SESSION_ACCESS.get)
        VECTOR_STORE.pop(oldest, None)
        SESSION_ACCESS.pop(oldest, None)
        HASH_STORE.pop(oldest, None)
        HASH_BY_DOC.pop(oldest, None)
        _evict_session_store(oldest)
        _cleanup_session_files(oldest)
    SESSION_ACCESS[sid] = now
    manifest = _manifest_path(sid)
    current_mtime = None
    if manifest.exists():
        try:
            current_mtime = manifest.stat().st_mtime
        except OSError:
            pass

    store = VECTOR_STORE.get(sid)
    if store is None:
        _load_session_manifest(sid)
        store = _make_store(sid)
        VECTOR_STORE[sid] = store
        if current_mtime is not None:
            _MANIFEST_MTIMES[sid] = current_mtime
    else:
        last_mtime = _MANIFEST_MTIMES.get(sid)
        if current_mtime is not None and (last_mtime is None or current_mtime > last_mtime):
            _load_session_manifest(sid)
            try:
                store.load()
            except Exception as e:
                logger.warning("Failed to reload store mirror during worker sync for %s: %s", sid, e)
            _MANIFEST_MTIMES[sid] = current_mtime
        elif current_mtime is None and last_mtime is not None:
            SESSION_FILES.pop(sid, None)
            _MANIFEST_MTIMES.pop(sid, None)
            store.chunks, store.vectors = [], []
            store._tfidf_index_cache = None
    return store


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _save_upload_to(upload: UploadFile, destination: Path) -> None:
    """Stream an UploadFile to disk — the FastAPI equivalent of Werkzeug's
    FileStorage.save(). Copied in chunks rather than via .read(), so a large
    (but still within MAX_CONTENT_LENGTH) upload never has to sit in memory
    in its entirety."""
    upload.file.seek(0)
    with open(destination, "wb") as out:
        shutil.copyfileobj(upload.file, out, length=1024 * 1024)
    upload.file.seek(0)


# ─────────────────────────────────────────────────────────────────────────────
#  Text Extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf_pages(filepath: str) -> list[dict]:
    """Extract text from each page of a PDF. Returns list of {page, text}."""
    pages = []
    try:
        doc = fitz.open(filepath)
    except Exception:
        logger.info("Could not open PDF %s (corrupt, encrypted, or not a real PDF)", filepath, exc_info=True)
        return []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            pages.append({"page": page_num, "text": text})
    doc.close()
    return pages


def extract_txt_pages(filepath: str) -> list[dict]:
    """Treat a .txt or .md file as logical sections separated by blank lines."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    sections = [s.strip() for s in re.split(r"\n{3,}", text) if s.strip()]
    if not sections:
        sections = [text]
    return [{"page": i + 1, "text": s} for i, s in enumerate(sections)]


_OCR_PROMPT = (
    "Extract all text, table content, diagram descriptions, titles, bullet points, "
    "and key information from this image into clean, structured Markdown text."
)


def _gemini_vision_ocr(img_b64: str, mime_type: str, filename: str) -> str:
    """Vision OCR via Gemini's generateContent API (multimodal inline_data)."""
    url = f"{GEMINI_URL}/models/{GEMINI_VISION_MODEL}:generateContent?key={GEMINI_API_KEY}"
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"inline_data": {"mime_type": mime_type, "data": img_b64}},
                    {"text": _OCR_PROMPT},
                ]
            }
        ],
        "generationConfig": {"temperature": 0}
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=90) as resp:
        data = json.loads(resp.read().decode("utf-8"))
    extracted_text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
    logger.info("🖼️ Gemini Vision OCR extracted %d characters from image %s", len(extracted_text), filename)
    return extracted_text


def _ollama_vision_ocr(img_b64: str, filename: str) -> str:
    """Vision OCR via a locally-running Ollama vision model (llava/moondream/
    llama3.2-vision/etc.) using the native /api/generate endpoint's
    `images` field. No API key, no rate limit, runs on your own machine.
    Expect noticeably weaker dense-text OCR than Gemini's — worth
    checking output quality on a real scan before relying on it."""
    url = f"{OLLAMA_URL}/api/generate"
    payload = {
        "model": OLLAMA_VISION_MODEL,
        "prompt": _OCR_PROMPT,
        "images": [img_b64],
        "stream": False,
        "options": {"temperature": 0},
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=180) as resp:
        data = json.loads(resp.read().decode("utf-8"))
    extracted_text = (data.get("response") or "").strip()
    logger.info("🖼️ Ollama (%s) Vision OCR extracted %d characters from image %s",
                OLLAMA_VISION_MODEL, len(extracted_text), filename)
    return extracted_text


def extract_image_pages(filepath: str) -> list[dict]:
    """Extract text and visual content from images via Vision OCR — Gemini
    or a local Ollama vision model, per VISION_BACKEND."""
    ext = filepath.rsplit(".", 1)[-1].lower()
    mime_type = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "webp": "image/webp",
        "bmp": "image/bmp",
        "tiff": "image/tiff",
        "gif": "image/gif",
    }.get(ext, "image/jpeg")

    filename = Path(filepath).name

    if VISION_BACKEND == "gemini" and not GEMINI_API_KEY:
        logger.warning("⚠️ Gemini API Key missing for Image Vision OCR: %s", filename)
        return [{"page": 1, "text": f"Document Image: {filename}\nNote: Configure GEMINI_API_KEY to enable full Vision OCR extraction."}]

    try:
        with open(filepath, "rb") as img_file:
            img_b64 = base64.b64encode(img_file.read()).decode("utf-8")

        if VISION_BACKEND == "ollama":
            extracted_text = _ollama_vision_ocr(img_b64, filename)
        else:
            extracted_text = _gemini_vision_ocr(img_b64, mime_type, filename)

        return [{"page": 1, "text": f"# Image OCR Content: {filename}\n\n{extracted_text}"}]
    except Exception as exc:
        backend_label = "Ollama" if VISION_BACKEND == "ollama" else "Gemini"
        logger.error("❌ %s Vision OCR failed for %s: %s", backend_label, filename, exc, exc_info=True)
        return [{"page": 1, "text": f"Image Document: {filename}\n(Error during image text extraction)"}]


def extract_docx_pages(filepath: str) -> list[dict]:
    """Extract text and tables from Word (.docx / .doc) documents."""
    try:
        if docx is None:
            raise ImportError("python-docx package not installed")
        doc = docx.Document(filepath)
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        for t in doc.tables:
            for row in t.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(f"| {row_text} |")
        text = "\n\n".join(full_text)
        sections = [s.strip() for s in re.split(r"\n{3,}", text) if s.strip()] or [text]
        return [{"page": i + 1, "text": s} for i, s in enumerate(sections)]
    except Exception as exc:
        logger.info("Falling back to plain text read for Word document %s: %s", filepath, exc)
        return extract_txt_pages(filepath)


def extract_data_pages(filepath: str, ext: str) -> list[dict]:
    """Extract CSV, TSV, JSON, XML, YAML data files as structured Markdown."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read().strip()

    if ext in ("csv", "tsv"):
        lines = content.splitlines()
        delimiter = "\t" if ext == "tsv" else ","
        formatted = []
        for line in lines:
            parts = line.split(delimiter)
            formatted.append(" | ".join(p.strip() for p in parts))
        text = f"```table\n" + "\n".join(formatted) + "\n```"
    elif ext in ("json", "yaml", "yml", "xml"):
        text = f"```{ext}\n{content}\n```"
    else:
        text = content

    return [{"page": 1, "text": text}]


def extract_code_pages(filepath: str, ext: str) -> list[dict]:
    """Extract source code files wrapped in code blocks."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        code = f.read().strip()
    wrapped = f"```{ext}\n{code}\n```"
    return [{"page": 1, "text": wrapped}]


def extract_document_pages(filepath: str, ext: str) -> list[dict]:
    """Universal text extractor supporting PDF, TXT, MD, DOCX, CSV/JSON, Code, and Images (Vision OCR)."""
    ext = ext.lower()
    if ext == "pdf":
        return extract_pdf_pages(filepath)
    elif ext in ("png", "jpg", "jpeg", "webp", "bmp", "tiff", "gif"):
        return extract_image_pages(filepath)
    elif ext in ("docx", "doc"):
        return extract_docx_pages(filepath)
    elif ext in ("csv", "tsv", "json", "yaml", "yml", "xml"):
        return extract_data_pages(filepath, ext)
    elif ext in ("py", "js", "ts", "jsx", "tsx", "html", "css", "c", "cpp", "h", "hpp", "java", "go", "rs", "php", "sql", "sh", "bat", "ps1"):
        return extract_code_pages(filepath, ext)
    else:
        return extract_txt_pages(filepath)


# ─────────────────────────────────────────────────────────────────────────────
#  Web Page Loading
# ─────────────────────────────────────────────────────────────────────────────

_IGNORE_TAGS = {"script", "style", "noscript", "svg", "canvas", "nav", "header", "footer"}


class _TextExtractor(html.parser.HTMLParser):
    def __init__(self):
        super().__init__()
        self.skip_depth = 0
        self.in_title = False
        self.title = None
        self.parts: list[str] = []

    def handle_starttag(self, tag, attrs):
        if tag in _IGNORE_TAGS:
            self.skip_depth += 1
        if tag == "title":
            self.in_title = True
        if tag in ("p", "div", "li", "h1", "h2", "h3", "h4", "tr", "br", "blockquote"):
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in _IGNORE_TAGS:
            self.skip_depth = max(0, self.skip_depth - 1)
        if tag == "title":
            self.in_title = False
        if tag in ("p", "div", "li", "h1", "h2", "h3", "h4", "blockquote"):
            self.parts.append("\n")

    def handle_data(self, data):
        if self.skip_depth:
            return
        if self.in_title:
            self.title = data.strip()
            return
        self.parts.append(data)

    def get_text(self) -> str:
        text = "".join(self.parts)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n+", "\n\n", text)
        return text.strip()


MAX_URL_FETCH_BYTES = 10 * 1024 * 1024  # 10 MB maximum web page download limit
URL_FETCH_TIMEOUT = 15  # 15 seconds per fetch hop


def _is_public_ip(ip_str: str) -> bool:
    try:
        ip = ipaddress.ip_address(ip_str)
    except ValueError:
        return False
    return not (
        ip.is_private or ip.is_loopback or ip.is_link_local
        or ip.is_multicast or ip.is_reserved or ip.is_unspecified
    )


def _validate_url_is_public(url: str) -> None:
    """
    Raise ValueError if `url` resolves to a private/loopback/link-local/
    reserved address or unauthorized internal port.
    """
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("Only http:// and https:// URLs are allowed")
    hostname = parsed.hostname
    if not hostname:
        raise ValueError("URL has no hostname")

    h_lower = hostname.lower()
    if (h_lower in ("localhost", "0.0.0.0", "127.0.0.1", "::1", "169.254.169.254")
            or h_lower.endswith(".local") or h_lower.endswith(".internal")):
        raise ValueError(f"URL hostname '{hostname}' resolves to a non-public/internal address — refusing to fetch")

    if parsed.port and parsed.port not in (80, 443, 8080, 8443):
        raise ValueError(f"Port {parsed.port} is not allowed for web document ingestion")

    try:
        infos = socket.getaddrinfo(hostname, None)
    except socket.gaierror as exc:
        raise ValueError(f"Could not resolve hostname '{hostname}': {exc}") from exc

    if not infos:
        raise ValueError(f"Could not resolve hostname '{hostname}' to an IP address")

    for family, _, _, _, sockaddr in infos:
        ip_str = sockaddr[0]
        if not _is_public_ip(ip_str):
            raise ValueError(
                f"URL '{hostname}' resolves to a non-public address ({ip_str}) — refusing to fetch"
            )


class _NoAutoRedirect(urllib.request.HTTPRedirectHandler):
    """Disables urllib's automatic redirect-following. Returning None here
    makes urlopen raise HTTPError with the redirect's status code instead
    of silently following it — see fetch_web_page() for why that matters."""
    def redirect_request(self, req, fp, code, msg, headers, newurl):
        return None


def fetch_web_page(url: str, max_redirects: int = 5) -> tuple[str, str]:
    """
    Fetch a URL and return (title, cleaned text) with SSRF checks on every redirect hop,
    response size limiting, and strict connection timeouts.
    """
    opener = urllib.request.build_opener(_NoAutoRedirect)
    current = url
    raw = None

    for _ in range(max_redirects + 1):
        _validate_url_is_public(current)
        req = urllib.request.Request(
            current,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 AskMyDocs/1.0"},
        )
        try:
            with opener.open(req, timeout=URL_FETCH_TIMEOUT) as resp:
                raw_bytes = resp.read(MAX_URL_FETCH_BYTES + 1)
                if len(raw_bytes) > MAX_URL_FETCH_BYTES:
                    raise ValueError(f"Page exceeded {MAX_URL_FETCH_BYTES // (1024 * 1024)} MB size limit")
                raw = raw_bytes.decode("utf-8", errors="ignore")
            break
        except urllib.error.HTTPError as e:
            if e.code in (301, 302, 303, 307, 308):
                location = e.headers.get("Location")
                if not location:
                    raise ValueError(f"Redirect ({e.code}) with no Location header") from e
                current = urllib.parse.urljoin(current, location)
                continue
            raise
    else:
        raise ValueError(f"Too many redirects (>{max_redirects})")

    if raw is None:
        raise ValueError("Too many redirects")

    parser = _TextExtractor()
    parser.feed(raw)
    title = parser.title or urllib.parse.urlparse(current).netloc
    text = parser.get_text()
    return title, text


# ─────────────────────────────────────────────────────────────────────────────
#  Fixed-Size Chunking
# ─────────────────────────────────────────────────────────────────────────────

def fixed_chunk(doc_info: dict, pages: list[dict], chunk_size: int,
                overlap_ratio: float = 0.2) -> list[dict]:
    """Split document text into overlapping fixed-size word chunks."""
    overlap = max(1, int(chunk_size * overlap_ratio))
    step = max(1, chunk_size - overlap)
    chunks = []
    chunk_index = 0

    for page_data in pages:
        words = page_data["text"].split()
        i = 0
        while i < len(words):
            end = min(i + chunk_size, len(words))
            text = " ".join(words[i:end])
            if text.strip():
                chunks.append({
                    "id": f"{doc_info['doc_id']}::p{page_data['page']}::c{chunk_index}",
                    "doc_id": doc_info["doc_id"],
                    "filename": doc_info["filename"],
                    "page": page_data["page"],
                    "section": None,
                    "block_type": "paragraph",
                    "text": text,
                    "chunk_index": chunk_index,
                    "method": f"{chunk_size}",
                })
                chunk_index += 1
            i += step

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
#  Structured Chunking
# ─────────────────────────────────────────────────────────────────────────────

STRUCT_MIN_WORDS = 40
STRUCT_MAX_WORDS = 500  # was 300 — real sections with a multi-point lettered
# list (e.g. a 7-item duties list plus a closing sentence) were getting
# split mid-section at the old limit, and only ONE fragment would end up
# among the top retrieval results — stranding the rest of that same
# section's content in a separate chunk that never got shown at all.
# Confirmed directly: a real section's source-card preview showed only
# its closing sentence, with the actual substantive list content missing
# entirely from the answer. 500 gives real-world sections like this
# enough room to survive as one coherent, fully-retrievable chunk.


_NUMBERED_HEADING_RE = re.compile(r"^\d+(\.\d+)*\.?\s+[A-Z][A-Za-z].*$")
_ALLCAPS_HEADING_RE = re.compile(r"^[A-Z][A-Z0-9 ,&'/\-]{2,70}$")


def _looks_like_plain_heading(stripped: str) -> bool:
    """
    Detects section headings in real-world documents that were never
    written in Markdown — numbered ("2.1 Duties, Rights and Obligations
    of GESCI") or ALL-CAPS ("INTRODUCTION", "OFFENCES") titles, the
    overwhelmingly common style in Word/PDF-extracted policy manuals,
    contracts, and handbooks.

    Constrained to short, title-like lines specifically so it does NOT
    misfire on numbered body sentences like "6.6. These records will be
    updated as required by law..." — a real heading is a few words with
    no trailing sentence punctuation; a numbered clause is a full
    sentence and will fail the word-count/punctuation checks below.
    """
    if not stripped or len(stripped) > 90:
        return False
    words = stripped.split()
    if not (1 <= len(words) <= 12):
        return False
    if _NUMBERED_HEADING_RE.match(stripped) and not stripped.endswith((".", ";", ",")):
        return True
    if _ALLCAPS_HEADING_RE.match(stripped) and any(c.isalpha() for c in stripped):
        return True
    return False


def parse_blocks(text: str, base_section: str = "Overview") -> list[dict]:
    """
    Parse plain text (or Markdown) into semantic blocks.
    Detects: headings (## style), code fences (```), table rows (|…|),
    and paragraph boundaries (blank lines).
    """
    lines = text.split("\n")
    blocks = []
    current_section = base_section
    buffer: list[str] = []
    block_type = "paragraph"
    in_code = False

    def flush():
        joined = " ".join(buffer).replace("  ", " ").strip()
        if not joined:
            buffer.clear()
            return
        words = len(joined.split())
        if block_type in ("code", "table") or words > 4:
            blocks.append({
                "type": block_type,
                "section": current_section,
                "text": joined,
            })
        buffer.clear()

    for line in lines:
        stripped = line.strip()

        # Code fence toggle
        if stripped.startswith("```"):
            if not in_code:
                flush()
                in_code = True
                block_type = "code"
            else:
                flush()
                in_code = False
                block_type = "paragraph"
            continue

        if in_code:
            if stripped:
                buffer.append(stripped)
            continue

        # Heading detection — Markdown style, or plain numbered/ALL-CAPS
        # (see _looks_like_plain_heading for why the latter is needed)
        heading_match = re.match(r"^#{1,4}\s+(.*)", stripped)
        if heading_match:
            flush()
            current_section = heading_match.group(1).strip()
            current_section = re.sub(r"\*\*(.+?)\*\*", r"\1", current_section)
            current_section = re.sub(r"`(.+?)`", r"\1", current_section)
            block_type = "paragraph"
            continue

        if _looks_like_plain_heading(stripped):
            flush()
            current_section = stripped
            block_type = "paragraph"
            continue

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            if block_type != "table":
                flush()
                block_type = "table"
            if not re.match(r"^[\|\s\-:]+$", stripped):
                cell_text = re.sub(r"\|", " ", stripped).strip()
                buffer.append(cell_text)
            continue

        # Blank line = paragraph boundary
        if not stripped:
            flush()
            block_type = "paragraph"
            continue

        # Normal line — strip Markdown noise
        cleaned = stripped
        cleaned = re.sub(r"^[-*>]\s+", "", cleaned)
        cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
        cleaned = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", cleaned)
        if cleaned:
            buffer.append(cleaned)

    flush()
    return blocks


SENTENCE_ABBREV = {
    "e.g", "i.e", "etc", "vs", "dr", "mr", "mrs", "ms", "prof", "st",
    "inc", "ltd", "co", "no", "vol", "fig", "al", "et", "approx",
}

_ABBREV_RE = re.compile(
    r"\b(" + "|".join(re.escape(a) for a in SENTENCE_ABBREV) + r")\.", re.IGNORECASE)


def split_sentences(text: str) -> list[str]:
    """
    Split text into sentences without breaking on common abbreviations
    (e.g. "e.g.", "Dr.") or decimal numbers ("1.5").
    """
    # Protect abbreviations and decimals so their periods aren't sentence endings
    protected = _ABBREV_RE.sub(lambda m: m.group(0).replace(".", "\x00"), text)
    protected = re.sub(r"\b(\d+)\.(\d+)\b",
                       lambda m: m.group(1) + "\x00" + m.group(2), protected)

    # Sentence ends: punctuation followed by whitespace + capital/CJK, or end-of-string
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z\u4e00-\u9fff])|(?<=[.!?])$", protected)

    sentences = []
    for part in parts:
        part = part.strip().replace("\x00", ".")
        if part:
            sentences.append(part)
    return sentences


def split_by_sentences(block: dict, max_words: int) -> list[dict]:
    result = []
    buf: list[str] = []
    count = 0

    for sent in split_sentences(block["text"]):
        w = len(sent.split())
        if count + w > max_words and buf:
            result.append({**block, "text": " ".join(buf)})
            buf = [sent]
            count = w
        else:
            buf.append(sent)
            count += w

    if buf:
        result.append({**block, "text": " ".join(buf)})
    return result


def _make_chunk(doc_info: dict, block: dict, index: int) -> dict:
    pages = block.get("pages") or [block.get("page", 1)]
    section = block.get("section")
    body_text = block["text"]
    # Prepend the section heading to the chunk's SEARCHABLE text — without
    # this, a query that closely matches a section's own title wording
    # (e.g. "Rights and Obligations of Staff" matching a section literally
    # titled "Duties, Rights and Obligations of Staff") scores badly,
    # because the heading's words were previously stripped into metadata
    # only and never appeared anywhere TF-IDF or embeddings actually look.
    # "Overview" is the fallback default when no real heading was ever
    # detected, so it's excluded here — it isn't a real title worth
    # matching against.
    if section and section != "Overview":
        searchable_text = f"{section}. {body_text}"
    else:
        searchable_text = body_text
    return {
        "id": f"{doc_info['doc_id']}::c{index}",
        "doc_id": doc_info["doc_id"],
        "filename": doc_info["filename"],
        "page": pages[0],
        "page_end": pages[-1] if len(pages) > 1 else None,
        "section": section,
        "block_type": block.get("type", "paragraph"),
        "text": searchable_text,
        "chunk_index": index,
        "method": "structured",
    }


def structured_chunk(doc_info: dict, pages: list[dict]) -> list[dict]:
    """
    Structured chunking pipeline:
      1. Parse each page into semantic blocks
      2. Merge tiny blocks (< STRUCT_MIN_WORDS)
      3. Split huge blocks (> STRUCT_MAX_WORDS) at sentence boundaries
      4. Attach metadata: section title, block type, page, filename
    """
    chunks = []
    chunk_index = 0
    merge_buffer: dict | None = None

    def flush_merge():
        nonlocal merge_buffer, chunk_index
        if merge_buffer is not None:
            chunks.append(_make_chunk(doc_info, merge_buffer, chunk_index))
            chunk_index += 1
            merge_buffer = None

    for page_data in pages:
        page_num = page_data["page"]
        blocks = parse_blocks(page_data["text"])

        for block in blocks:
            words = len(block["text"].split())

            if words < STRUCT_MIN_WORDS:
                if merge_buffer is None:
                    merge_buffer = {**block, "page": page_num, "pages": [page_num]}
                else:
                    if merge_buffer["pages"][-1] != page_num:
                        merge_buffer["pages"].append(page_num)
                    merge_buffer["text"] += " " + block["text"]
                    # Track the section of whichever block was merged in
                    # MOST RECENTLY, not just whichever started the merge —
                    # otherwise a chunk combining "10 SEPARATION FROM
                    # SERVICE" (short intro) + "10.1 Resignation" (short
                    # detail) keeps citing the intro's heading even though
                    # the more specific 10.1 content is what actually
                    # answers a question about resignation notice periods.
                    if block.get("section"):
                        merge_buffer["section"] = block["section"]
                    if len(merge_buffer["text"].split()) >= STRUCT_MIN_WORDS:
                        flush_merge()
                continue

            flush_merge()

            if words > STRUCT_MAX_WORDS:
                for sub in split_by_sentences(block, STRUCT_MAX_WORDS):
                    sub["page"] = page_num
                    sub["pages"] = [page_num]
                    chunks.append(_make_chunk(doc_info, sub, chunk_index))
                    chunk_index += 1
            else:
                block["page"] = page_num
                block["pages"] = [page_num]
                chunks.append(_make_chunk(doc_info, block, chunk_index))
                chunk_index += 1

    flush_merge()

    return chunks


def chunk_text(doc_info: dict, pages: list[dict], chunk_mode: str) -> list[dict]:
    if chunk_mode == "structured":
        return structured_chunk(doc_info, pages)
    size = int(chunk_mode) if chunk_mode.isdigit() else 256
    return fixed_chunk(doc_info, pages, size)


# ─────────────────────────────────────────────────────────────────────────────
#  Pure-Python TF-IDF Retrieval (offline fallback)
# ─────────────────────────────────────────────────────────────────────────────

STOPWORDS = {
    "the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "do", "does", "did", "will", "would", "could",
    "should", "may", "might", "shall", "can", "to", "of", "in", "for",
    "on", "with", "at", "by", "from", "up", "about", "into", "through",
    "and", "but", "or", "if", "while", "when", "where", "who", "which",
    "that", "this", "these", "those", "it", "its", "all", "each", "both",
    "more", "most", "other", "some", "no", "not", "only", "same", "than",
    "too", "very", "just", "your", "you", "they", "we", "he", "she", "i",
    "my", "their", "our", "his", "her", "also", "how", "what", "so", "as",
}

CONFIDENCE_THRESHOLD = 0.05
# CONFIDENCE_THRESHOLD above only filters which chunks are even considered
# candidates — it's deliberately low so search_chunks() doesn't miss
# anything. It is NOT a "should we actually answer" gate: a single
# incidentally-shared word (e.g. both the query and an unrelated document
# containing the word "policy") can clear 0.05 easily, producing a
# confident-looking answer with source citations for a question that has
# nothing to do with the document. TFIDF_MIN_SCORE below is the real gate
# — checked via validate_context() in /ask, the same way EMBED_MIN_SCORE
# already gates the embeddings path.
TFIDF_MIN_SCORE = float(os.environ.get("TFIDF_MIN_SCORE", "0.15"))


def tokenize(text: str) -> list[str]:
    """Lowercase, extract word tokens (incl. Unicode/CJK), remove stopwords."""
    tokens = _simple_tokenizer.findall(text.lower())
    return [t for t in tokens if len(t) > 2 and t not in STOPWORDS]


def compute_tf(tokens: list[str]) -> dict[str, float]:
    if not tokens:
        return {}
    counts = Counter(tokens)
    total = len(tokens)
    return {t: c / total for t, c in counts.items()}


def compute_idf(corpus_tokens: list[list[str]]) -> dict[str, float]:
    n = len(corpus_tokens)
    doc_freq: dict[str, int] = {}
    for tokens in corpus_tokens:
        for t in set(tokens):
            doc_freq[t] = doc_freq.get(t, 0) + 1
    return {t: math.log((1 + n) / (1 + df)) + 1.0 for t, df in doc_freq.items()}


def tfidf_vector(tf: dict[str, float], idf: dict[str, float]) -> dict[str, float]:
    return {t: tf_val * idf.get(t, 1.0) for t, tf_val in tf.items()}


def cosine_sim(vec_a: dict[str, float], vec_b: dict[str, float]) -> float:
    common = set(vec_a) & set(vec_b)
    if not common:
        return 0.0
    dot = sum(vec_a[t] * vec_b[t] for t in common)
    norm_a = math.sqrt(sum(v * v for v in vec_a.values()))
    norm_b = math.sqrt(sum(v * v for v in vec_b.values()))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


BM25_K1 = float(os.environ.get("BM25_K1", "1.5"))  # term-frequency saturation —
# higher = a repeated word keeps adding score for longer before flattening out
BM25_B = float(os.environ.get("BM25_B", "0.75"))  # document-length normalization —
# 0 = ignore length entirely (like TF-IDF effectively does), 1 = fully
# penalize longer chunks for their length


def build_index(chunks: list[dict]) -> dict:
    corpus_tokens = [tokenize(c["text"]) for c in chunks]
    doc_lengths = [len(tokens) for tokens in corpus_tokens]
    avg_doc_length = sum(doc_lengths) / len(doc_lengths) if doc_lengths else 0.0
    # BM25's own IDF variant (log((N - n + 0.5) / (n + 0.5) + 1)) is
    # deliberately different from compute_idf() above — it's shaped to
    # never go negative even when a term appears in more than half the
    # corpus, which compute_idf()'s plain formula can do.
    n = len(corpus_tokens)
    doc_freq: dict[str, int] = {}
    for tokens in corpus_tokens:
        for t in set(tokens):
            doc_freq[t] = doc_freq.get(t, 0) + 1
    bm25_idf = {t: math.log((n - df + 0.5) / (df + 0.5) + 1) for t, df in doc_freq.items()}
    return {
        "corpus_tokens": corpus_tokens,
        "idf": compute_idf(corpus_tokens),
        "doc_lengths": doc_lengths,
        "avg_doc_length": avg_doc_length,
        "bm25_idf": bm25_idf,
    }


def bm25_score(query_tokens: list[str], doc_tokens: list[str], doc_length: int,
              avg_doc_length: float, bm25_idf: dict[str, float]) -> float:
    """
    Real BM25, not TF-IDF. The two things this adds that plain TF-IDF
    cosine similarity doesn't have:

    1. Term-frequency SATURATION — TF-IDF's raw term-frequency component
       scales roughly linearly, so a word appearing 20 times in a chunk
       scores dramatically higher than appearing once, even past the
       point where more repetitions tell you anything more about
       relevance. BM25's (k1+1)*f / (f+k1) term flattens out — it keeps
       increasing but with steeply diminishing returns, matching the
       intuition that "mentioned it 20 times" isn't 20x more relevant
       than "mentioned it once."

    2. Document-length NORMALIZATION — TF-IDF has no notion of chunk
       length at all, so a long chunk that happens to contain a query
       word once looks identical (per that word) to a short, focused
       chunk containing it once, even though the short chunk is
       intuitively a much stronger, more concentrated match. BM25's
       (1 - b + b * doc_length/avg_doc_length) term in the denominator
       penalizes matches diluted across a long chunk relative to the
       corpus's average chunk length.
    """
    if avg_doc_length == 0:
        return 0.0
    doc_tf = Counter(doc_tokens)
    score = 0.0
    length_norm = 1 - BM25_B + BM25_B * (doc_length / avg_doc_length)
    for t in query_tokens:
        f = doc_tf.get(t, 0)
        if f == 0:
            continue
        idf = bm25_idf.get(t, 0.0)
        score += idf * (f * (BM25_K1 + 1)) / (f + BM25_K1 * length_norm)
    return score


def bm25_scores_for_corpus(query: str, index: dict) -> list[float]:
    """BM25 score of `query` against every chunk in `index`, in corpus order."""
    query_tokens = tokenize(query)
    if not query_tokens:
        return [0.0] * len(index["corpus_tokens"])
    return [
        bm25_score(query_tokens, doc_tokens, index["doc_lengths"][i],
                  index["avg_doc_length"], index["bm25_idf"])
        for i, doc_tokens in enumerate(index["corpus_tokens"])
    ]


def search_chunks(query: str, chunks: list[dict], index: dict,
                  top_k: int = 4) -> list[dict]:
    """Real BM25 search — offline fallback (was plain TF-IDF cosine
    similarity; see bm25_score()'s docstring for exactly what changed
    and why)."""
    if not chunks:
        return []

    query_tokens = tokenize(query)
    if not query_tokens:
        return []

    raw_scores = bm25_scores_for_corpus(query, index)
    # BM25 scores are unbounded (unlike TF-IDF cosine's [0,1] range) and
    # their scale depends on corpus size/IDF, so CONFIDENCE_THRESHOLD
    # (calibrated for the old TF-IDF cosine scale) doesn't directly
    # apply. Normalize by the max score in this result set so downstream
    # thresholds (TFIDF_MIN_SCORE, the shared-token safeguard) keep
    # working against a comparable [0,1]-ish scale.
    max_score = max(raw_scores) if raw_scores else 0.0
    scored = []
    for chunk, score in zip(chunks, raw_scores):
        normalized = (score / max_score) if max_score > 0 else 0.0
        if normalized >= CONFIDENCE_THRESHOLD:
            scored.append({**chunk, "score": normalized})

    scored.sort(key=lambda x: -x["score"])
    return scored[:top_k]


RRF_K = int(os.environ.get("RRF_K", "10"))  # NOT the textbook default (60) —
# that's calibrated for web-scale search with thousands of candidates.
# Tested directly for a document-sized corpus (a few hundred chunks):
# k=60 gives a perfect rank-0-in-both match ~1.0 but a plausible FALSE
# POSITIVE (moderate rank ~15-20 in both methods) ~0.78 — nearly as high,
# destroying the threshold's ability to tell them apart. k=10 spreads
# these to ~1.0 vs ~0.39-0.42, leaving real room for EMBED_MIN_SCORE to
# discriminate.


def reciprocal_rank_fusion(store: "VectorStore", query: str, top_k: int = 5) -> list[dict]:
    """
    Combines embedding-based ranking and real BM25-based ranking (not
    TF-IDF — see bm25_score() for exactly what that adds) via
    Reciprocal Rank Fusion, replacing the old weighted-average blend
    (alpha*embed + (1-alpha)*tfidf) that hybrid_search() used to do.

    Why weighted averaging kept failing: it's scale-sensitive. A chunk
    that's genuinely the #1 best match by embedding similarity can still
    have a modest-looking raw cosine score (e.g. 0.3 instead of 0.7) for
    reasons having nothing to do with correctness — different embedding
    models, different document styles, short/generic-sounding queries
    all shift the RAW SCORE distribution without changing which chunk is
    actually most relevant. Averaging that modest raw score against
    TF-IDF's raw score can pull an obviously-correct match below a fixed
    threshold — repeatedly observed in production: multiple different
    exact-section-title queries scored 0.12-0.19 against a 0.15-0.30
    threshold despite correctly ranking #1 in their own method.

    RRF sidesteps this entirely: it only asks "what RANK did this chunk
    get in each method's own ordering," never "how big was the raw
    score." A chunk ranked #1 by embeddings AND #1 by TF-IDF scores near
    the maximum regardless of what the underlying raw numbers happened
    to be. This is the standard, named technique for combining ranked
    lists from different retrieval methods (Cormack et al., 2009).

    The final score is normalized to roughly a real number in [0, 1]
    (1.0 = ranked #1 in both methods) purely so it remains compatible
    with the existing validate_context()/EMBED_MIN_SCORE threshold logic
    downstream — the actual fusion math above the normalization step is
    unmodified textbook RRF.
    """
    if not store.chunks:
        return []
    n = len(store.chunks)

    embed_scores = (store.query_scores(embed_text(query))
                    if store.vectors and len(store.vectors) == len(store.chunks)
                    else [0.0] * n)
    embed_order = sorted(range(n), key=lambda i: -embed_scores[i])
    embed_rank = {chunk_i: rank for rank, chunk_i in enumerate(embed_order)}

    index = store.get_tfidf_index()
    keyword_scores = bm25_scores_for_corpus(query, index)
    keyword_order = sorted(range(n), key=lambda i: -keyword_scores[i])
    keyword_rank = {chunk_i: rank for rank, chunk_i in enumerate(keyword_order)}

    max_possible = 2.0 / (RRF_K + 1)  # rank 0 (best) in both methods
    rrf_scores = [
        (1.0 / (RRF_K + embed_rank[i] + 1) + 1.0 / (RRF_K + keyword_rank[i] + 1)) / max_possible
        for i in range(n)
    ]

    ranked = sorted(range(n), key=lambda i: -rrf_scores[i])[:top_k]
    return [
        {**store.chunks[i], "score": rrf_scores[i],
         "embed_score": embed_scores[i], "keyword_score": keyword_scores[i]}
        for i in ranked
    ]


def hybrid_search(store: "VectorStore", query: str, top_k: int = 5,
                  alpha: float = HYBRID_ALPHA) -> list[dict]:
    """
    Rank chunks by a weighted blend of embedding similarity and TF-IDF
    similarity, instead of embeddings alone.

    Kept for RETRIEVAL_MODE=hybrid-legacy / A-B comparison against
    reciprocal_rank_fusion() above, which is now the default (see
    RETRIEVAL_MODE handling in /ask). Superseded as the default because
    it's scale-sensitive in a way RRF isn't — see that function's
    docstring for the specific, repeatedly-observed failure mode.
    """
    if not store.chunks:
        return []

    embed_scores = (store.query_scores(embed_text(query))
                    if store.vectors and len(store.vectors) == len(store.chunks)
                    else [0.0] * len(store.chunks))

    index = store.get_tfidf_index()
    q_vec = tfidf_vector(compute_tf(tokenize(query)), index["idf"])
    tfidf_scores = [
        cosine_sim(q_vec, tfidf_vector(compute_tf(tokens), index["idf"]))
        for tokens in index["corpus_tokens"]
    ]

    combined = [alpha * e + (1 - alpha) * t for e, t in zip(embed_scores, tfidf_scores)]
    ranked = sorted(range(len(combined)), key=lambda i: -combined[i])[:top_k]

    return [
        {**store.chunks[i], "score": combined[i],
         "embed_score": embed_scores[i], "tfidf_score": tfidf_scores[i]}
        for i in ranked
    ]


def synthesize_answer(query: str, results: list[dict]) -> dict:
    """Template-based answer for the offline fallback path."""
    if not results:
        return {
            "found": False,
            "answer": (
                "I don't know — I couldn't find a relevant answer in the uploaded "
                "documents. Try uploading more documents or rephrasing your question."
            ),
            "sources": [],
        }

    seen: set[str] = set()
    sources = []
    all_text = " ".join(r["text"] for r in results)

    for r in results:
        # Dedupe by chunk ID, not page number — two genuinely distinct,
        # relevant sections often share the same page in a real document
        # (e.g. two subsections on one dense page), and page-based
        # dedup was silently dropping legitimate additional sources
        # whenever that happened, even though the answer text below
        # draws from all of them.
        key = r["id"]
        if key not in seen:
            seen.add(key)
            sources.append(r)

    query_words = set(tokenize(query))
    sentences = split_sentences(all_text)

    scored_sents = []
    for s in sentences:
        s = s.strip()
        if len(s) < 20:
            continue
        s_words = set(tokenize(s))
        overlap = len(query_words & s_words)
        scored_sents.append((overlap, s))

    scored_sents.sort(key=lambda x: -x[0])

    seen_sents: set[str] = set()
    top_sents = []
    for overlap, s in scored_sents:
        # Stop once we run out of GENUINELY relevant sentences — sentences
        # are sorted by overlap descending, so once overlap hits 0 every
        # remaining sentence is equally irrelevant. Previously this loop
        # kept going regardless, padding a simple question's answer with
        # up to 4 completely unrelated sentences just to force the count
        # to exactly 5, even when only 1 sentence actually answered it.
        if overlap <= 0:
            break
        if s not in seen_sents:
            seen_sents.add(s)
            top_sents.append(s)
        # Still cap the upper end — a genuinely complex question where
        # many sentences share query terms shouldn't run away unbounded.
        if len(top_sents) >= 8:
            break

    answer = " ".join(top_sents) if top_sents else results[0]["text"][:500] + "…"

    return {"found": True, "answer": answer, "sources": sources}


# ─────────────────────────────────────────────────────────────────────────────
#  Validation + LLM-grounded generation
# ─────────────────────────────────────────────────────────────────────────────

_DONT_KNOW_MARKERS = (
    "i don't know", "i do not know", "cannot answer", "can't answer",
    "not enough information", "doesn't contain", "does not contain",
    "not available in the", "no information",
)


def _stem_lite(token: str) -> str:
    """Crude suffix-stripping so simple variations (discover/discovers,
    obligation/obligations) count as the same word for the shared-token
    safeguard below — without pulling in a real stemmer or changing the
    main tokenize()/TF-IDF pipeline used everywhere else in scoring."""
    for suffix in ("ing", "es", "ed", "s"):
        if len(token) > len(suffix) + 3 and token.endswith(suffix):
            return token[: -len(suffix)]
    return token


def validate_context(results: list[dict], min_score: float, query: str = None,
                     rerank_score: float | None = None) -> bool:
    """
    Reject retrieval when nothing clears the similarity bar, OR when
    there's no good evidence the top result is actually relevant.
    """
    if not results:
        return False
    top_score = results[0].get("score", 0.0)
    if top_score < min_score:
        return False
    if rerank_score is not None:
        return rerank_score >= RERANK_MIN_RELEVANCE
    if query is not None:
        # Generic intent/question framing words that should not artificially inflate the required token count
        META_QUESTION_WORDS = {
            "meaning", "mean", "definition", "define", "explain", "explanation",
            "describe", "description", "overview", "detail", "details", "tell",
            "show", "give", "list", "state", "clarify", "understand", "concept",
            "purpose", "reason", "procedure", "process", "rule", "rules"
        }
        query_tokens = set(tokenize(query))
        if query_tokens:
            chunk_tokens = set(tokenize(results[0]["text"]))
            query_stems = {_stem_lite(t) for t in query_tokens}
            chunk_stems = {_stem_lite(t) for t in chunk_tokens}
            shared = query_stems & chunk_stems

            # Content stems excluding conversational meta-question words
            content_stems = {s for s in query_stems if s not in META_QUESTION_WORDS}

            # If top retrieval confidence is very strong (e.g. >= 0.75 or 1.0 in RRF),
            # 1 shared stem is plenty to confirm it's on-topic and avoid false rejections.
            if top_score >= 0.75:
                required = 1
            elif content_stems:
                required = min(2, len(content_stems))
            else:
                required = min(2, len(query_stems))

            if len(shared) < required:
                return False
    return True


def _is_dont_know(answer: str | None) -> bool:
    """
    Detects a genuine refusal, not just the incidental presence of a
    refusal-like phrase somewhere in a real answer.
    """
    if not answer or not answer.strip():
        return True
    a = answer.lower().strip()
    if a.rstrip(".").strip() in ("i don't know", "i do not know"):
        return True
    word_count = len(a.split())
    return word_count <= 12 and any(m in a for m in _DONT_KNOW_MARKERS)


RERANK_ENABLED = os.environ.get("RERANK_ENABLED", "false").lower() == "true"
RERANK_TOP_N = int(os.environ.get("RERANK_TOP_N", "8"))  # how many RRF candidates to rerank
RERANK_MIN_RELEVANCE = float(os.environ.get("RERANK_MIN_RELEVANCE", "5"))  # 0-10 scale;
# the reranker's own relevance judgment must clear this to answer at all


def rerank_with_llm(query: str, results: list[dict]) -> tuple[list[dict], float | None]:
    """
    A second-pass reranker over RRF's already-narrowed candidate list —
    functionally the same PURPOSE as a cross-encoder (Cohere Rerank /
    BGE-Reranker, as named in the syllabus): look at a small candidate
    set more carefully than a first-pass ranker can, and reorder it.

    Why this isn't literally a cross-encoder: a real one needs either a
    paid API (Cohere) or a local model (BGE-Reranker, via
    sentence-transformers + torch) — a large new dependency this
    project's stack doesn't otherwise need. This uses the LLM call
    infrastructure already in place instead: send the query and each
    candidate's text to the LLM, ask for a 0-10 relevance score per
    candidate, and reorder by that score. Same place in the pipeline,
    same goal, different (already-available) mechanism — not a silent
    swap presented as something it isn't.

    Also returns the TOP candidate's relevance score (0-10, or None if
    reranking didn't run). This exists specifically so /ask can use it
    in place of the shared-token safeguard when reranking is active: a
    pure lexical-overlap check rejects genuinely correct but
    low-overlap answers (e.g. "who issued this?" vs "...issued by the
    Financial Conduct Authority" shares almost no words), which a real
    relevance judgment from the LLM correctly recognizes as a match.
    The token-overlap check remains the fallback when reranking is OFF,
    since there'd otherwise be no smarter judge available at all.

    Fails open: if the LLM call errors for any reason, returns the
    original RRF order unchanged and a None score (caller falls back
    to the token-overlap check in that case) rather than breaking
    retrieval.
    """
    if not results or not _chat_configured():
        return results, None

    candidates = results[:RERANK_TOP_N]
    numbered = "\n\n".join(f"[{i}] {c['text'][:600]}" for i, c in enumerate(candidates))
    system = RERANK_SYSTEM_PROMPT
    user = f"Question: {query}\n\nCandidates:\n{numbered}"
    try:
        raw = _chat_call(system, user, temperature=0, max_tokens=300)
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
        scored = json.loads(raw)
        indices = sorted(item["index"] for item in scored)
        if indices != list(range(len(candidates))):
            raise ValueError("LLM did not score every candidate exactly once")
        scored.sort(key=lambda item: -item["score"])
        reranked = [candidates[item["index"]] for item in scored]
        top_score = float(scored[0]["score"])
        return reranked + results[RERANK_TOP_N:], top_score  # anything beyond RERANK_TOP_N stays untouched, in original order
    except Exception:
        logger.warning("LLM reranking failed — falling back to the original RRF order.",
                       exc_info=True)
        return results, None


QUERY_REWRITE_ENABLED = os.environ.get("QUERY_REWRITE_ENABLED", "false").lower() == "true"

RERANK_SYSTEM_PROMPT = register_prompt(RERANK_PROMPT_VERSION, (
    "You score search results for relevance to a question. Given a "
    "question and numbered candidate excerpts, respond with ONLY a "
    "JSON array of objects, one per candidate, each with \"index\" "
    "and \"score\" (0-10, where 10 means the excerpt directly and "
    "fully answers the question, 0 means completely unrelated — "
    "judge genuine relevance, not just shared words). Order the "
    "array from highest score to lowest. No other text — just the "
    "JSON array, e.g. [{\"index\":2,\"score\":9},{\"index\":0,\"score\":3}]."
))

REWRITE_SYSTEM_PROMPT = register_prompt(REWRITE_PROMPT_VERSION, (
    "Rewrite the user's question into a short, keyword-rich search "
    "query optimized for retrieving relevant passages from a "
    "document — not a natural-language answer, not a rephrased "
    "question, just the core searchable terms. Respond with ONLY "
    "the rewritten query, nothing else."
))

QA_SYSTEM_PROMPT = register_prompt(QA_PROMPT_VERSION, (
    "You are a grounded question-answering assistant. Answer using ONLY the "
    "document excerpts provided. Cite every fact with its source number like "
    "[1] or [2]. If the excerpts do not contain enough information to answer "
    "the question, reply exactly with: \"I don't know.\" Do not use outside "
    "knowledge. Match your answer's length to what the question actually "
    "needs — a short, direct sentence or two for a simple factual question, "
    "but a longer, complete answer (including every item, if the source "
    "material itself lists several, such as a numbered or lettered list) "
    "for a question that calls for it. Never omit relevant details from the "
    "excerpts just to keep the answer short; completeness for the specific "
    "question asked matters more than brevity."
))


def rewrite_query(query: str) -> str:
    """
    Rewrites a messy/conversational question into a cleaner search query
    BEFORE retrieval — e.g. "so like what happens if someone just doesn't
    show up to work for a while" -> "abandonment of post policy".

    Only the rewritten query is used for retrieval (embeddings + BM25).
    The ORIGINAL question is still what's shown in the UI and what gets
    sent to generate_answer() for the final answer — rewriting only
    changes what we search FOR, never what we tell the user we searched,
    and never what the model is actually asked to answer.

    Fails open: returns the original query unchanged on any error.
    """
    if not query.strip() or not _chat_configured():
        return query
    system = REWRITE_SYSTEM_PROMPT
    try:
        rewritten = _chat_call(system, query, temperature=0, max_tokens=60).strip('"')
        return rewritten if rewritten else query
    except Exception:
        logger.warning("Query rewriting failed — using the original question as-is.",
                       exc_info=True)
        return query


def generate_answer(query: str, results: list[dict], temperature: float = 0.0) -> str:
    """Ask the LLM to produce a grounded answer with [N] source citations."""
    system = QA_SYSTEM_PROMPT
    user = build_qa_user_prompt(query, results)

    try:
        return _chat_call(system, user, temperature=temperature)
    except Exception:
        return ""


def build_qa_user_prompt(query: str, results: list[dict]) -> str:
    """Rebuild the exact user-turn text generate_answer() sends, from a
    results list alone. Shared by generate_answer() and replay so the two
    can never silently drift apart."""
    context_blocks = []
    for i, r in enumerate(results, start=1):
        loc = f"{r['filename']} (page {r['page']})"
        if r.get("section"):
            loc += f", section: {r['section']}"
        context_blocks.append(f"[{i}] {loc}\n{r['text']}")
    return "DOCUMENTS:\n\n" + "\n\n".join(context_blocks) + f"\n\nQUESTION: {query}\n\nANSWER:"


# ─────────────────────────────────────────────────────────────────────────────
#  Request / response schemas
# ─────────────────────────────────────────────────────────────────────────────
# These exist for two concrete reasons: real request validation at the edge
# (instead of hand-rolled `.get()` chains inside every handler) and an
# accurate, self-updating OpenAPI schema at /docs. Response models are used
# only where the payload shape is genuinely fixed — the endpoints whose
# bodies vary by branch (/ask, /upload, /eval/run, ...) keep returning plain
# dicts so the wire format stays byte-for-byte what the frontend already
# expects.


class _LenientModel(BaseModel):
    """Base for request bodies: unknown keys are ignored rather than rejected,
    matching the previous `request.get_json(silent=True) or {}` tolerance."""

    model_config = ConfigDict(extra="ignore")


class UploadCancelRequest(_LenientModel):
    upload_id: str | None = None


class LoadUrlRequest(_LenientModel):
    url: str | None = None
    chunk_mode: str | None = None


class AskRequest(_LenientModel):
    query: str = ""
    chunk_mode: str | None = None
    top_k: int | None = None
    temperature: float | None = None


class RemoveRequest(_LenientModel):
    doc_id: str | None = None


class EvalQuestion(_LenientModel):
    id: str | None = None
    question: str | None = None
    expected: str | None = None
    expected_doc: str | None = None
    expected_section: str | None = None


class EvalRunRequest(_LenientModel):
    questions: list[EvalQuestion] = []
    top_k: int | None = None
    k: int = 3
    strategy_filter: str | None = None
    chunk_mode: str | None = None
    presets: list[str] | None = None
    modes: list[str] | None = None  # legacy alias for `presets`


class OkResponse(BaseModel):
    ok: bool = True


class ClearResponse(BaseModel):
    ok: bool = True
    warning: str | None = None


class RemoveResponse(BaseModel):
    ok: bool = True
    removed_chunks: int = 0
    warning: str | None = None


class StatusDocument(BaseModel):
    filename: str
    doc_id: str
    chunk_count: int
    method: str
    openable: bool


class StatusResponse(BaseModel):
    total_chunks: int
    documents: list[StatusDocument]
    methods: list[str]
    mode: str
    vector_backend: str


class HealthzResponse(BaseModel):
    status: str
    embeddings_configured: bool
    chat_configured: bool
    chat_backend: str
    embeddings_backend: str
    retrieval_mode: str
    vector_backend: str
    active_sessions: int


class ReadyzResponse(BaseModel):
    ready: bool
    checks: dict[str, bool]


class TracesResponse(BaseModel):
    count: int
    trace_ids: list[str]


# ─────────────────────────────────────────────────────────────────────────────
#  Session dependencies
# ─────────────────────────────────────────────────────────────────────────────
# Flask's `session` was an implicit request-global; Starlette's lives on the
# Request. Rather than repeat the same three lines at the top of every
# handler, the two access patterns the app actually needs are expressed as
# dependencies: "give me a session, creating one if this is a first visit"
# and "there must already be one, else 400".


class NoActiveSessionError(Exception):
    """Raised by require_session_id() when the caller has no session cookie."""


@app.exception_handler(NoActiveSessionError)
async def _no_active_session_handler(request: Request, exc: NoActiveSessionError) -> JSONResponse:
    return JSONResponse({"error": "No active session"}, status_code=400)


@app.exception_handler(RequestValidationError)
async def _validation_error_handler(request: Request, exc: RequestValidationError) -> JSONResponse:
    """Every error this API returns carries an `error` string — that is what
    frontend/src/services/api.ts's handleResponse() reads. FastAPI's default 422 body is
    `{"detail": [...]}` instead, which the frontend would render as a bare
    "Request failed with status 422". This keeps the envelope consistent
    while still exposing the structured `detail` for API clients and /docs.
    """
    problems = "; ".join(
        f"{'.'.join(str(p) for p in err.get('loc', ())[1:]) or 'body'}: {err.get('msg', 'invalid')}"
        for err in exc.errors()
    )
    return JSONResponse(
        {"error": f"Invalid request: {problems}", "detail": jsonable_encoder(exc.errors())},
        status_code=422,
    )


def ensure_session_id(request: Request) -> str:
    """Session id for this browser, minting one on first contact."""
    sid = request.session.get("session_id")
    if not sid:
        sid = str(uuid.uuid4())
        request.session["session_id"] = sid
    return sid


def require_session_id(request: Request) -> str:
    """Session id for this browser, or a 400 {"error": "No active session"}."""
    sid = request.session.get("session_id")
    if not sid:
        raise NoActiveSessionError()
    return sid


def optional_session_id(request: Request) -> str | None:
    """Session id if there is one, else None — for endpoints that answer
    differently (or with a different status code) when there's no session."""
    return request.session.get("session_id")


def require_session_store(sid: Annotated[str, Depends(require_session_id)]) -> VectorStore:
    """This session's loaded vector store (Qdrant- or memory-backed), or a
    400 if there is no session at all."""
    return _get_store(sid)


SessionId = Annotated[str, Depends(ensure_session_id)]
RequiredSessionId = Annotated[str, Depends(require_session_id)]
OptionalSessionId = Annotated[str | None, Depends(optional_session_id)]
RequiredStore = Annotated[VectorStore, Depends(require_session_store)]


# ─────────────────────────────────────────────────────────────────────────────
#  Routes
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/", include_in_schema=False)
def index(request: Request, sid: SessionId):
    # `sid` is unused in the body on purpose: depending on it is what mints
    # the session cookie for a first-time visitor.
    index_html = FRONTEND_DIST / "index.html"
    if index_html.exists():
        return FileResponse(str(index_html))
    return JSONResponse(
        {"message": "React frontend not built. Run 'npm run build' in the frontend/ directory."},
        status_code=503,
    )


def _index_into_store(sid: str, doc_info: dict, pages: list[dict],
                      chunk_mode: str, chunks_out: list[dict],
                      embedding_ok: bool) -> None:
    """Chunk pages and add them to the session's vector store."""
    store = _get_store(sid)
    new_chunks = chunk_text(doc_info, pages, chunk_mode)
    chunks_out.append(new_chunks)

    if embedding_ok:
        vectors = embed_texts([c["text"] for c in new_chunks])
    else:
        vectors = []
    store.add(new_chunks, vectors)


@app.post("/upload-cancel", response_model=OkResponse)
def upload_cancel(payload: UploadCancelRequest | None = Body(default=None)):
    """
    Signals that an in-flight /upload call (identified by the same
    upload_id the client sent with its FormData) should be cancelled.
    This just records the signal — /upload itself is what actually acts
    on it, checking between files and rolling back anything already
    added before it returns. If /upload has already finished and
    returned by the time this arrives, the signal is simply ignored
    (nothing left to cancel) and cleaned up by the TTL sweep.
    """
    _sweep_cancelled_uploads()
    upload_id = ((payload.upload_id if payload else None) or "").strip()
    if upload_id:
        CANCELLED_UPLOADS[upload_id] = time.time()
    return OkResponse()


@app.post("/upload")
def upload(
    sid: SessionId,
    files: list[UploadFile] = File(default=[]),
    chunk_mode: str = Form(default="structured"),
    upload_id: str = Form(default=""),
):
    upload_id = (upload_id or "").strip()
    _sweep_cancelled_uploads()

    if not files:
        return JSONResponse({"error": "No files provided"}, status_code=400)

    logger.info("📤 Upload request received: %d file(s) (chunk_mode: %s, session: %s)",
                len(files), chunk_mode, sid[:8])
    hashes = HASH_STORE.setdefault(sid, set())
    embedding_ok = _embeddings_configured()
    results = []
    pending: list[dict] = []

    RAGTracer.trace("INGESTION", 1, 5, "Upload Request Received", {
        "Files Count": len(files),
        "Chunk Strategy": chunk_mode,
        "Session ID": sid[:8],
        "Embeddings Enabled": embedding_ok,
    })

    for f in files:
        if not f or not allowed_file(f.filename or ""):
            logger.warning("⚠️ Unsupported file type uploaded: %s", getattr(f, "filename", "?"))
            results.append({"filename": getattr(f, "filename", "?") or "?",
                            "error": "Unsupported file type (allowed: PDF, TXT, MD)"})
            continue

        original_name = f.filename
        ext = (original_name.rsplit(".", 1)[1] if "." in original_name else "").lower()

        safe_name = secure_filename(original_name) or "document"
        if not safe_name.lower().endswith("." + ext):
            safe_name = f"{safe_name}.{ext}"
        filepath = UPLOAD_FOLDER / f"{uuid.uuid4()}_{safe_name}"

        dedupe_key = None
        try:
            _save_upload_to(f, filepath)
            with open(filepath, "rb") as fh:
                content_hash = hashlib.sha256(fh.read()).hexdigest()
            dedupe_key = (content_hash, chunk_mode)

            if dedupe_key in hashes:
                filepath.unlink(missing_ok=True)
                logger.info("ℹ️ Duplicate upload skipped for '%s' under mode '%s'", original_name, chunk_mode)
                results.append({"filename": original_name,
                                "error": f"Already indexed under the '{chunk_mode}' strategy — "
                                         f"pick a different chunking strategy to compare, or remove "
                                         f"the existing one first."})
                continue
            hashes.add(dedupe_key)

            doc_id = str(uuid.uuid4())[:8]
            doc_info = {"doc_id": doc_id, "filename": original_name}

            pages = (extract_pdf_pages(str(filepath)) if ext == "pdf"
                     else extract_txt_pages(str(filepath)))

            if not pages:
                filepath.unlink(missing_ok=True)
                hashes.discard(dedupe_key)
                reason = ("No extractable text (password-protected or scanned PDF?)"
                          if ext == "pdf" else "Empty file")
                logger.warning("⚠️ Text extraction returned empty for '%s': %s", original_name, reason)
                results.append({"filename": original_name, "error": reason})
                continue

            new_chunks = chunk_text(doc_info, pages, chunk_mode)
            if not new_chunks:
                filepath.unlink(missing_ok=True)
                hashes.discard(dedupe_key)
                logger.warning("⚠️ Chunking produced 0 chunks for '%s'", original_name)
                results.append({"filename": original_name, "error": "No chunks produced"})
                continue

            RAGTracer.trace("INGESTION", 2, 5, "Text Extraction & Chunking", {
                "File Name": original_name,
                "Format": ext.upper(),
                "Pages Extracted": len(pages),
                "Chunks Generated": len(new_chunks),
                "Doc ID": doc_id,
            })

            CHUNK_COUNTS.setdefault(sid, {})
            for mode in ("structured", "128", "256", "512"):
                CHUNK_COUNTS[sid][mode] = len(chunk_text(doc_info, pages, mode))

            pending.append({
                "filepath": filepath,
                "doc_id": doc_id,
                "filename": original_name,
                "ext": ext,
                "chunks": new_chunks,
                "hash": dedupe_key,
                "result": {
                    "filename": original_name,
                    "pages": len(pages),
                    "chunks": len(new_chunks),
                    "method": chunk_mode,
                },
            })
        except Exception as exc:
            filepath.unlink(missing_ok=True)
            if dedupe_key is not None:
                hashes.discard(dedupe_key)
            logger.error("❌ Exception during processing '%s': %s", original_name, exc, exc_info=True)
            results.append({"filename": original_name, "error": f"Failed to index: {exc}"})

    if not pending:
        return JSONResponse({"ok": False, "error": "No valid documents were indexed.",
                             "documents": results}, status_code=400)

    store = _get_store(sid)
    committed_doc_ids: list[str] = []  # tracks what THIS call actually added, for rollback below
    was_cancelled = False

    for item in pending:
        if upload_id and upload_id in CANCELLED_UPLOADS:
            # Stop processing any remaining files in this batch — no
            # point embedding/indexing more once cancellation was
            # requested. Whatever's already committed gets rolled back
            # below, after the loop.
            was_cancelled = True
            item["filepath"].unlink(missing_ok=True)
            hashes.discard(item["hash"])
            logger.warning("🚫 Upload cancelled by client for upload_id %s", upload_id)
            continue

        added_to_store = False
        stored_path = None
        try:
            embedding_degraded = False
            if VECTOR_BACKEND == "qdrant":
                if not embedding_ok:
                    raise ValueError(
                        "Qdrant vector backend requires embeddings. "
                        "Configure an embedding backend (e.g. GEMINI_API_KEY or OLLAMA) first."
                    )
                try:
                    vectors = embed_texts([c["text"] for c in item["chunks"]])
                except Exception as embed_exc:
                    raise ValueError(f"Embedding generation failed on Qdrant backend: {embed_exc}") from embed_exc
            else:
                # Memory/TF-IDF backend: can degrade gracefully if embedding fails
                if embedding_ok:
                    try:
                        vectors = embed_texts([c["text"] for c in item["chunks"]])
                    except Exception as embed_exc:
                        logger.warning(
                            "⚠️ Embedding failed for '%s', indexing without vectors "
                            "(keyword search only for this document): %s",
                            item["filename"], embed_exc,
                        )
                        vectors = []
                        embedding_degraded = True
                else:
                    vectors = []

            store.add(item["chunks"], vectors)
            added_to_store = True

            stored_path = UPLOAD_FOLDER / f"{sid}__{item['doc_id']}.{item['ext']}"
            item["filepath"].replace(stored_path)
            SESSION_FILES.setdefault(sid, {})[item["doc_id"]] = {
                "path": stored_path,
                "name": item["filename"],
            }
            _save_session_manifest(sid)

            # Remember which hash belongs to this doc_id so /remove can
            # release it later (see remove_doc()) — otherwise re-uploading
            # the same file after removing it is wrongly flagged duplicate.
            HASH_BY_DOC.setdefault(sid, {})[item["doc_id"]] = item["hash"]

            item["result"]["doc_id"] = item["doc_id"]
            item["result"]["openable"] = True
            if embedding_degraded:
                item["result"]["warning"] = (
                    "Indexed with keyword search only — embeddings failed "
                    "(check EMBED_BACKEND config / Ollama server). Semantic "
                    "search won't work for this document until re-uploaded."
                )
            results.append(item["result"])
            committed_doc_ids.append(item["doc_id"])
            RAGTracer.trace("INGESTION", 5, 5, "Store Commit Completed", {
                "File Name": item["filename"],
                "Doc ID": item["doc_id"],
                "Chunks Indexed": len(item["chunks"]),
                "Vector Backend": VECTOR_BACKEND,
                "Total Session Chunks": len(store.chunks),
            })
            logger.info("✅ Indexed '%s' (%d chunks, vectors: %s, total store: %d chunks)",
                        item["filename"], len(item["chunks"]), "yes" if bool(vectors) else "no", len(store.chunks))
        except Exception as exc:
            # Transaction rollback for this document
            cleanup_complete = True
            cleanup_error = None
            orphan_rec = None
            if added_to_store:
                try:
                    store.remove_doc(item["doc_id"])
                    _resolve_orphaned_doc(sid, item["doc_id"])
                except Exception as rb_exc:
                    cleanup_complete = False
                    cleanup_error = f"Vector store rollback failed: {rb_exc}"
                    logger.error(
                        "❌ CRITICAL: Failed to rollback vector store for doc %s ('%s'): %s. Vector may be orphaned.",
                        item["doc_id"], item["filename"], rb_exc, exc_info=True,
                    )
                    orphan_rec = _record_orphaned_doc(
                        sid=sid,
                        doc_id=item["doc_id"],
                        filename=item["filename"],
                        error=str(rb_exc),
                        stored_path=stored_path or item["filepath"],
                    )
            if stored_path and stored_path.exists():
                stored_path.unlink(missing_ok=True)
            item["filepath"].unlink(missing_ok=True)
            hashes.discard(item["hash"])
            SESSION_FILES.get(sid, {}).pop(item["doc_id"], None)
            HASH_BY_DOC.get(sid, {}).pop(item["doc_id"], None)
            try:
                _save_session_manifest(sid)
            except Exception:
                pass
            logger.error("❌ Failed to index '%s': %s", item["filename"], exc, exc_info=True)
            doc_res = {
                "filename": item["filename"],
                "error": f"Failed to index: {exc}",
                "cleanup_complete": cleanup_complete,
            }
            if not cleanup_complete:
                doc_res["cleanup_error"] = cleanup_error
                doc_res["doc_id"] = item["doc_id"]
                if orphan_rec and orphan_rec.get("reconciliation_persistence_failed"):
                    doc_res["reconciliation_persistence_failed"] = True
            results.append(doc_res)

    # Final check: even if cancellation only arrived after every file had
    # already finished committing, honor it anyway. Attempt to stop/rollback
    # this upload's work. Backend deletion can fail, so incomplete cleanup
    # is explicitly reported and reconciliation metadata is retained when possible.
    if upload_id and upload_id in CANCELLED_UPLOADS:
        was_cancelled = True
    if was_cancelled:
        logger.warning("🚫 Rolling back %d committed doc(s) due to upload cancellation", len(committed_doc_ids))
        failed_rollbacks = []
        for doc_id in committed_doc_ids:
            rollback_ok = False
            try:
                store.remove_doc(doc_id)
                rollback_ok = True
                _resolve_orphaned_doc(sid, doc_id)
            except Exception as rb_err:
                logger.error("❌ Failed to remove doc %s from store during cancellation: %s", doc_id, rb_err, exc_info=True)
                file_info = SESSION_FILES.get(sid, {}).get(doc_id)
                orphan_rec = _record_orphaned_doc(
                    sid=sid,
                    doc_id=doc_id,
                    filename=file_info.get("name") if file_info else "unknown",
                    error=str(rb_err),
                    stored_path=file_info.get("path") if file_info else None,
                )
                failed_rollbacks.append({
                    "doc_id": doc_id,
                    "error": str(rb_err),
                    "reconciliation_persistence_failed": bool(orphan_rec.get("reconciliation_persistence_failed")),
                })

            if rollback_ok:
                info = SESSION_FILES.get(sid, {}).pop(doc_id, None)
                if info:
                    info["path"].unlink(missing_ok=True)
                doc_hash = HASH_BY_DOC.get(sid, {}).pop(doc_id, None)
                if doc_hash is not None:
                    hashes.discard(doc_hash)
            else:
                # Retain metadata in SESSION_FILES flagged as orphan so reconciliation is possible
                if sid in SESSION_FILES and doc_id in SESSION_FILES[sid]:
                    SESSION_FILES[sid][doc_id]["orphan"] = True
                    SESSION_FILES[sid][doc_id]["rollback_error"] = str(failed_rollbacks[-1]["error"])

        _save_session_manifest(sid)
        CANCELLED_UPLOADS.pop(upload_id, None)

        if failed_rollbacks:
            return {
                "ok": False,
                "cancelled": True,
                "cleanup_complete": False,
                "error": f"Upload cancelled but cleanup was incomplete: failed to remove {len(failed_rollbacks)} document(s) from vector store.",
                "failed_cleanup": failed_rollbacks,
                "documents": [],
            }
        else:
            return {
                "ok": False,
                "cancelled": True,
                "cleanup_complete": True,
                "error": "Upload cancelled — nothing was indexed.",
                "documents": [],
            }

    return {"ok": True, "documents": results, "total_chunks": len(store.chunks),
            "chunk_comparison": CHUNK_COUNTS.get(sid, {})}


@app.post("/load-url")
def load_url(sid: SessionId, payload: LoadUrlRequest | None = Body(default=None)):
    payload = payload or LoadUrlRequest()
    url = (payload.url or "").strip()
    chunk_mode = payload.chunk_mode or "structured"

    if not url:
        return JSONResponse({"error": "Empty URL"}, status_code=400)

    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        return JSONResponse({"error": "URL must start with http:// or https://"}, status_code=400)

    embedding_ok = _embeddings_configured()

    try:
        title, text = fetch_web_page(url)
    except Exception as exc:
        return JSONResponse({"error": f"Failed to fetch URL: {exc}"}, status_code=400)

    if len(text.split()) < 20:
        return JSONResponse({"error": "Page returned too little text to index."}, status_code=400)

    doc_id = str(uuid.uuid4())[:8]
    doc_info = {"doc_id": doc_id, "filename": title[:80] or parsed.netloc}
    pages = [{"page": 1, "text": text}]

    new_chunks = chunk_text(doc_info, pages, chunk_mode)
    if not new_chunks:
        return JSONResponse({"error": "No chunks produced from this page."}, status_code=400)

    store = _get_store(sid)
    if VECTOR_BACKEND == "qdrant":
        if not embedding_ok:
            return JSONResponse({
                "error": "Qdrant vector backend requires embeddings. Configure an embedding backend first."
            }, status_code=503)
        try:
            vectors = embed_texts([c["text"] for c in new_chunks])
            store.add(new_chunks, vectors)
        except Exception as exc:
            return JSONResponse({"error": f"Embedding/indexing failed on Qdrant backend: {exc}"},
                                status_code=503)
    else:
        if embedding_ok:
            try:
                vectors = embed_texts([c["text"] for c in new_chunks])
                store.add(new_chunks, vectors)
            except Exception as exc:
                return JSONResponse({"error": f"Embedding failed: {exc}"}, status_code=500)
        else:
            store.add(new_chunks, [])

    CHUNK_COUNTS.setdefault(sid, {})
    for mode in ("structured", "128", "256", "512"):
        CHUNK_COUNTS[sid][mode] = len(chunk_text(doc_info, pages, mode))

    return {"ok": True, "documents": [{
        "filename": doc_info["filename"],
        "doc_id": doc_id,
        "openable": False,
        "pages": 1,
        "chunks": len(new_chunks),
        "method": chunk_mode,
    }], "total_chunks": len(store.chunks),
        "chunk_comparison": CHUNK_COUNTS.get(sid, {})}


@app.get("/favicon.ico", include_in_schema=False)
def favicon():
    """Serves the project's favicon."""
    for path in (
        FRONTEND_DIST / "favicon.ico",
        FRONTEND_DIST / "rag.svg",
    ):
        if path.exists():
            media_type = "image/x-icon" if path.suffix == ".ico" else "image/svg+xml"
            return FileResponse(path, media_type=media_type)
    return JSONResponse({"error": "Favicon not found"}, status_code=404)


@app.get("/file/{doc_id}", include_in_schema=False)
def serve_file(request: Request, doc_id: str, store: RequiredStore):
    """View an uploaded document in the React Single Page Application."""
    sid = request.session["session_id"]
    info = SESSION_FILES.get(sid, {}).get(doc_id)
    doc_chunks = [c for c in store.chunks if c["doc_id"] == doc_id]

    if not info and not doc_chunks:
        return JSONResponse({"error": "File not found or no longer available"}, status_code=404)

    index_html = FRONTEND_DIST / "index.html"
    if index_html.exists():
        return FileResponse(str(index_html))

    return JSONResponse(
        {"message": "React frontend not built. Run 'npm run build' in the frontend/ directory."},
        status_code=503,
    )


@app.get("/file/{doc_id}/raw")
def serve_file_raw(doc_id: str, sid: RequiredSessionId):
    """Stream the raw PDF bytes (used by the viewer page's embedded viewer)."""
    info = SESSION_FILES.get(sid, {}).get(doc_id)
    if not info or not info["path"].exists():
        return JSONResponse({"error": "File not found or no longer available"}, status_code=404)
    # content_disposition_type="inline" is what Flask's as_attachment=False did.
    # FileResponse defaults to "attachment" when a filename is given, which
    # would make the browser download the PDF instead of rendering it in the
    # viewer's <embed>.
    return FileResponse(
        info["path"], filename=info["name"], content_disposition_type="inline"
    )


@app.get("/file/{doc_id}/pages")
def serve_file_pages(doc_id: str, sid: RequiredSessionId, store: RequiredStore):
    """Return the extracted pages (for the text viewer) of an uploaded document."""
    info = SESSION_FILES.get(sid, {}).get(doc_id)
    doc_chunks = [c for c in store.chunks if c["doc_id"] == doc_id]

    if not info and not doc_chunks:
        return JSONResponse({"error": "File not found or no longer available"}, status_code=404)

    if info and info["path"].exists():
        name = info["name"]
        ext = (name.rsplit(".", 1)[1] if "." in name else "").lower()
        try:
            pages = extract_document_pages(str(info["path"]), ext)
        except Exception as exc:
            return JSONResponse({"error": f"Could not read document: {exc}"}, status_code=500)
    else:
        # Reconstruct text pages from vector store chunks!
        name = doc_chunks[0]["filename"] if doc_chunks else "Document"
        ext = "txt"
        combined_text = "\n\n".join(c["text"] for c in doc_chunks)
        pages = [{"page": 1, "text": combined_text}]

    return {
        "filename": name,
        "ext": ext,
        "total_pages": len(pages),
        "pages": [{"num": i + 1, "text": p.get("text", "")} for i, p in enumerate(pages)],
    }


@app.post("/ask")
def ask(sid: OptionalSessionId, payload: AskRequest | None = Body(default=None)):
    if not sid:
        return JSONResponse({"error": "No documents uploaded yet"}, status_code=400)

    payload = payload or AskRequest()
    query = (payload.query or "").strip()
    method_filter = (payload.chunk_mode or "").strip() or None

    # Dynamic TOP_K and TEMPERATURE from request payload (fallback to env/defaults).
    # Pydantic has already coerced/validated the types; these just clamp to the
    # supported range rather than rejecting an out-of-range value outright.
    top_k = max(1, min(20, payload.top_k)) if payload.top_k is not None else TOP_K
    temperature = max(0.0, min(1.0, payload.temperature)) if payload.temperature is not None else 0.0

    trace_id = str(uuid.uuid4())
    _t0 = time.time()

    def log_ask_trace(*, retrieval_mode, retrieved, valid_context, model,
                       temperature, prompt_version, raw_output, answer,
                       found, rerank_score=None, extra=None):
        """Writes one durable, replayable /ask trace record. Redaction runs
        here — the ONLY place a record is built — so no exit path from
        /ask can accidentally skip it."""
        record = {
            "trace_id": trace_id,
            "session_id_hash": hashlib.sha256(sid.encode()).hexdigest()[:16],
            "question": redact(query),
            "search_query": redact(search_query) if search_query != query else None,
            "chunk_mode_filter": method_filter,
            "retrieval_mode": retrieval_mode,
            "top_k": top_k,
            "embed_min_score": EMBED_MIN_SCORE,
            "rerank_enabled": RERANK_ENABLED,
            "rerank_score": rerank_score,
            "retrieved": redact_deep([
                {
                    "chunk_id": r.get("id"),
                    "doc_id": r.get("doc_id"),
                    "filename": r.get("filename"),
                    "page": r.get("page"),
                    "section": r.get("section"),
                    "score": round(r["score"], 4) if r.get("score") is not None else None,
                    "text": r.get("text"),
                }
                for r in (retrieved or [])
            ]),
            "valid_context": valid_context,
            "model": model,
            "temperature": temperature,
            "prompt_version": prompt_version,
            "raw_output": redact(raw_output),
            "answer": redact(answer),
            "found": found,
            "latency_ms": round((time.time() - _t0) * 1000, 1),
        }
        if extra:
            record.update(redact_deep(extra))
        TRACES.log(record)

    store = _get_store(sid)

    RAGTracer.trace("RETRIEVAL", 1, 6, "Question Received", {
        "User Question": query,
        "Session ID": sid[:8],
        "Strategy Filter": method_filter or "All",
        "Total Session Chunks": len(store.chunks),
        "Top-K": top_k,
        "Temperature": temperature,
    })

    search_query = rewrite_query(query) if QUERY_REWRITE_ENABLED else query
    if QUERY_REWRITE_ENABLED and search_query != query:
        RAGTracer.trace("RETRIEVAL", 2, 6, "Query Rewriting", {
            "Original Question": query,
            "Rewritten Search Query": search_query,
        })

    def annotate_openable(results):
        files = SESSION_FILES.get(sid, {})
        for r in results:
            r["openable"] = r["doc_id"] in files
        return results
    if not store.chunks:
        return {
            "found": False,
            "answer": "No documents have been uploaded yet. Please upload a PDF, text file, or web page first.",
            "sources": [],
            "top_k": top_k,
            "temperature": temperature,
        }

    active_store = store
    if method_filter:
        active_store = store.filtered_by_method(method_filter)
        if not active_store.chunks:
            return {
                "found": False,
                "answer": (f"No documents are indexed under the '{method_filter}' chunking "
                          f"strategy yet. Upload one under that strategy first, or ask "
                          f"without a strategy filter to search everything indexed."),
                "sources": [],
                "top_k": top_k,
                "temperature": temperature,
            }

    # Path 1: embeddings + LLM (if configured and vectors exist)
    if _embeddings_configured() and _chat_configured() and active_store.vectors and len(active_store.vectors) == len(active_store.chunks):
        try:
            if RETRIEVAL_MODE == "hybrid-legacy":
                raw_results = hybrid_search(active_store, search_query, top_k=top_k)
            elif RETRIEVAL_MODE == "embed":
                q_vec = embed_text(search_query)
                raw_results = active_store.query(q_vec, top_k=top_k, min_score=0.0)
            else:  # "hybrid" (default) — Reciprocal Rank Fusion
                raw_results = reciprocal_rank_fusion(active_store, search_query, top_k=top_k)
            near_miss = raw_results[0] if raw_results else None
            results = [r for r in raw_results if r["score"] >= EMBED_MIN_SCORE]
            
            RAGTracer.trace("RETRIEVAL", 3, 6, "Hybrid Retrieval & RRF Fusion", {
                "Retrieval Mode": RETRIEVAL_MODE,
                "Top-K Requested": top_k,
                "Raw Candidates Returned": len(raw_results),
                "Above Min Threshold (" + str(EMBED_MIN_SCORE) + ")": len(results),
                "Top Match Score": f"{near_miss['score']:.4f}" if near_miss else "0.0000",
                "Top Source File": near_miss["filename"] if near_miss else "None",
            })

            rerank_score = None
            if RERANK_ENABLED and len(results) > 1:
                results, rerank_score = rerank_with_llm(query, results)
                RAGTracer.trace("RETRIEVAL", 4, 6, "LLM Reranker Evaluation", {
                    "Reranker Status": "Active",
                    "Candidates Evaluated": len(results),
                    "Top Relevance Score": f"{rerank_score:.1f}/10" if rerank_score else "N/A",
                })

            results = fit_to_token_budget(results, MAX_CONTEXT_TOKENS)
            valid = validate_context(results, EMBED_MIN_SCORE, query, rerank_score)
            
            RAGTracer.trace("RETRIEVAL", 5, 6, "Context Validation Gate", {
                "Threshold Gate": "PASSED" if valid else "FAILED",
                "Min Required Score": EMBED_MIN_SCORE,
                "Context Token Budget": f"{MAX_CONTEXT_TOKENS} tokens",
                "Valid Grounding Chunks": len(results),
            })

            if not valid:
                log_ask_trace(
                    retrieval_mode=RETRIEVAL_MODE, retrieved=raw_results,
                    valid_context=False, model=LLM_MODEL, temperature=temperature,
                    prompt_version=None, raw_output=None,
                    answer="I don't know — no document content matched your question closely enough.",
                    found=False, rerank_score=rerank_score,
                )
                return {
                    "found": False,
                    "answer": "I don't know — no document content matched your question closely enough.",
                    "sources": [],
                    "trace_id": trace_id,
                    "top_k": top_k,
                    "temperature": temperature,
                    "closest_match": ({
                        "filename": near_miss["filename"],
                        "section": near_miss.get("section"),
                        "page": near_miss.get("page"),
                        "score": round(near_miss["score"], 4),
                        "threshold": EMBED_MIN_SCORE,
                    } if near_miss else None),
                }
            answer = generate_answer(query, results, temperature=temperature)
            RAGTracer.trace("RETRIEVAL", 6, 6, "LLM Answer Generation", {
                "Model Identifier": LLM_MODEL,
                "Temperature": temperature,
                "Sources Grounded": len(results),
                "Answer Character Count": len(answer),
            })
            log_ask_trace(
                retrieval_mode=RETRIEVAL_MODE, retrieved=results,
                valid_context=True, model=LLM_MODEL, temperature=temperature,
                prompt_version=QA_PROMPT_VERSION, raw_output=answer,
                answer=answer or "I don't know.",
                found=not _is_dont_know(answer), rerank_score=rerank_score,
            )
            return {
                "found": not _is_dont_know(answer),
                "answer": answer or "I don't know.",
                "sources": annotate_openable(results),
                "trace_id": trace_id,
                "top_k": top_k,
                "temperature": temperature,
            }
        except RetrievalBackendError as exc:
            logger.error("❌ Vector database retrieval failed in /ask: %s", exc, exc_info=True)
            return JSONResponse({
                "error": f"Vector database retrieval failed: {exc}",
                "status": 503,
            }, status_code=503)
        except urllib.error.HTTPError as exc:
            reason = {
                400: "Bad request — check your API key format and model name.",
                401: "Invalid or expired API key — check GEMINI_API_KEY (if EMBED_BACKEND/VISION_BACKEND=gemini) and XAI_API_KEY (if CHAT_BACKEND=xai).",
                403: "Access forbidden — verify your Gemini API key has the required permissions.",
                429: "Rate limited by Gemini API — too many requests too quickly.",
            }.get(exc.code, f"Unexpected HTTP {exc.code} from Gemini API.")
            logger.warning(
                "Embeddings/LLM path failed for a query — falling back to TF-IDF-only. %s",
                reason, exc_info=True,
            )
            # fall through to TF-IDF
        except Exception:
            logger.warning(
                "Embeddings/LLM path failed for a query — falling back to TF-IDF-only. "
                "This usually means a network issue reaching Gemini API.",
                exc_info=True,
            )
            # fall through to TF-IDF

    # Path 2: offline TF-IDF fallback
    chunks = active_store.chunks
    index = active_store.get_tfidf_index()
    results = search_chunks(search_query, chunks, index, top_k=top_k)
    results = fit_to_token_budget(results, MAX_CONTEXT_TOKENS)
    if not validate_context(results, TFIDF_MIN_SCORE, query):
        near_miss = results[0] if results else None
        log_ask_trace(
            retrieval_mode="tfidf", retrieved=results, valid_context=False,
            model="tfidf-template", temperature=None, prompt_version=None,
            raw_output=None,
            answer="I don't know — no document content matched your question closely enough.",
            found=False,
        )
        return {
            "found": False,
            "answer": "I don't know — no document content matched your question closely enough.",
            "sources": [],
            "trace_id": trace_id,
            "top_k": top_k,
            "temperature": temperature,
            "closest_match": ({
                "filename": near_miss["filename"],
                "section": near_miss.get("section"),
                "page": near_miss.get("page"),
                "score": round(near_miss["score"], 4),
                "threshold": TFIDF_MIN_SCORE,
            } if near_miss else None),
        }
    resp = synthesize_answer(query, results)
    resp["sources"] = annotate_openable(resp.get("sources", []))
    resp["trace_id"] = trace_id
    resp["top_k"] = top_k
    resp["temperature"] = temperature
    log_ask_trace(
        retrieval_mode="tfidf", retrieved=results, valid_context=True,
        model="tfidf-template", temperature=None, prompt_version=None,
        raw_output=resp.get("answer"), answer=resp.get("answer"),
        found=resp.get("found", True),
    )
    return resp


@app.get("/traces", response_model=TracesResponse)
def list_traces(current_sid: OptionalSessionId):
    """List logged trace_ids for the current session."""
    if not current_sid:
        return TracesResponse(count=0, trace_ids=[])
    expected_hash = hashlib.sha256(current_sid.encode()).hexdigest()[:16]
    session_traces = [r["trace_id"] for r in TRACES.all() if r.get("session_id_hash") == expected_hash and "trace_id" in r]
    return TracesResponse(count=len(session_traces), trace_ids=session_traces)


@app.post("/replay/{trace_id}")
def replay_trace(trace_id: str, current_sid: OptionalSessionId):
    """Replay requirement#1: privacy-preserving replay reconstructing the
    generation from the durable prompt template version artifact, recorded model/temperature,
    and persisted redacted context snapshot."""
    if not current_sid:
        return JSONResponse({"error": "No active session"}, status_code=401)

    record = TRACES.get(trace_id)
    if not record:
        return JSONResponse({"error": f"No trace found for trace_id={trace_id}"}, status_code=404)

    expected_hash = hashlib.sha256(current_sid.encode()).hexdigest()[:16]
    if record.get("session_id_hash") != expected_hash:
        return JSONResponse({"error": "Unauthorized: trace belongs to another session"}, status_code=403)

    missing = []
    for field in ("prompt_version", "model", "retrieved", "raw_output"):
        if record.get(field) in (None, [], ""):
            missing.append(field)

    if record.get("retrieval_mode") == "tfidf" or not record.get("prompt_version"):
        return {
            "trace_id": trace_id,
            "replayable": False,
            "reason": (
                "This trace has no LLM prompt_version (either the TF-IDF "
                "offline path answered it, or the context-validation gate "
                "rejected it before any LLM call was made — both produce "
                "deterministic, template-based output with no LLM call)."
            ),
            "original": {
                "answer": record.get("answer"),
                "retrieved": record.get("retrieved"),
            },
            "fields_missing_from_trace": missing,
        }

    prompt_text = get_prompt(record["prompt_version"])
    if prompt_text is None:
        # If exact prompt artifact is missing, abort replay immediately without invoking LLM
        return JSONResponse({
            "trace_id": trace_id,
            "replayable": False,
            "reason": f"Durable prompt artifact for prompt_version {record['prompt_version']!r} is missing from registry. Exact prompt version required to replay.",
            "original": {
                "question": record.get("question"),
                "model": record.get("model"),
                "prompt_version": record.get("prompt_version"),
                "raw_output": record.get("raw_output"),
            },
            "replayed": None,
            "outputs_match_exactly": None,
            "fields_missing_from_trace": missing,
        }, status_code=400)

    results_for_prompt = [
        {
            "filename": r.get("filename"),
            "page": r.get("page"),
            "section": r.get("section"),
            "text": r.get("text") or "",
        }
        for r in record.get("retrieved", [])
    ]
    user_prompt = build_qa_user_prompt(record["question"], results_for_prompt)

    try:
        replayed_raw = _chat_call(
            prompt_text, user_prompt,
            temperature=record.get("temperature") or 0.0,
            model=record.get("model"),
        )
        replay_error = None
    except Exception as exc:
        replayed_raw = None
        replay_error = str(exc)

    return {
        "trace_id": trace_id,
        "replayable": True,
        "replay_type": "privacy_preserving",
        "original": {
            "question": record.get("question"),
            "model": record.get("model"),
            "prompt_version": record.get("prompt_version"),
            "raw_output": record.get("raw_output"),
        },
        "replayed": {
            "model": record.get("model"),
            "prompt_version": record.get("prompt_version"),
            "raw_output": replayed_raw,
            "error": replay_error,
        },
        "outputs_match_exactly": (replayed_raw == record.get("raw_output")) if replayed_raw is not None else None,
        "fields_missing_from_trace": missing,
    }


@app.post("/clear", response_model=ClearResponse, response_model_exclude_none=True)
def clear(sid: OptionalSessionId):
    backend_error = None
    if sid:
        # _get_store() reconnects to the real backend (Qdrant collection
        # or pickle file) if it isn't already cached in this process's
        # memory — using VECTOR_STORE.pop() directly here would silently
        # no-op after any restart/worker respawn, since it'd find nothing
        # to pop and never actually call .clear() on the real backend.
        store = _get_store(sid)
        store.clear()
        backend_error = getattr(store, "last_backend_error", None)
        if not backend_error:
            orphans = list(ORPHANED_DOCS.get(sid, []))
            for o in orphans:
                if o.get("doc_id"):
                    _resolve_orphaned_doc(sid, o["doc_id"])
        VECTOR_STORE.pop(sid, None)
        SESSION_ACCESS.pop(sid, None)
        HASH_STORE.pop(sid, None)
        HASH_BY_DOC.pop(sid, None)
        CHUNK_COUNTS.pop(sid, None)
        _cleanup_session_files(sid)
    resp = ClearResponse(ok=True)
    if backend_error:
        resp.warning = (f"Cleared locally, but the vector database delete "
                        f"failed: {backend_error}. The data may still exist "
                        f"in the backend.")
    return resp


@app.post("/remove", response_model=RemoveResponse, response_model_exclude_none=True)
def remove_doc(sid: RequiredSessionId, payload: RemoveRequest | None = Body(default=None)):
    """Remove a single document: its chunks, its retained file, its manifest
    entry, and its content hash (so the same file can be re-uploaded later
    without being wrongly flagged as a duplicate)."""
    doc_id = ((payload.doc_id if payload else None) or "").strip()
    if not doc_id:
        return JSONResponse({"error": "Missing doc_id"}, status_code=400)

    removed_chunks = 0
    store = _get_store(sid)
    if any(c["doc_id"] == doc_id for c in store.chunks):
        removed_chunks = store.remove_doc(doc_id)
        backend_error = getattr(store, "last_backend_error", None)
        if not backend_error:
            _resolve_orphaned_doc(sid, doc_id)

    files = SESSION_FILES.get(sid, {})
    info = files.pop(doc_id, None)
    if info:
        info["path"].unlink(missing_ok=True)
    _save_session_manifest(sid)

    # Forget this doc's content hash so a future re-upload of the same
    # file isn't rejected as a stale duplicate.
    doc_hash = HASH_BY_DOC.get(sid, {}).pop(doc_id, None)
    if doc_hash is not None:
        HASH_STORE.get(sid, set()).discard(doc_hash)

    resp = RemoveResponse(ok=True, removed_chunks=removed_chunks)
    backend_error = getattr(store, "last_backend_error", None)
    if backend_error:
        # Removed from the local view either way, but the real backend
        # (Qdrant) delete failed — surface this so it's not just a
        # server-log warning nobody sees. The document may reappear if
        # the session reloads from the backend later.
        resp.warning = (f"Removed locally, but the vector database delete "
                        f"failed: {backend_error}. The data may still exist "
                        f"in the backend.")
    return resp


def parse_qa_pairs(text: str) -> list[dict]:
    """
    Parses "Q: ...\\nA: ..." blocks out of raw text (from a PDF, or typed
    directly) into {question, expected} pairs for the eval page.

    Handles a question or answer spanning multiple lines by accumulating
    lines until the next "Q:"/"A:" marker or end of text — so a longer,
    wrapped question in the source PDF doesn't get cut short.
    """
    pairs = []
    current_q, current_a = None, None

    def flush():
        if current_q and current_a:
            pairs.append({"question": current_q.strip(), "expected": current_a.strip()})

    mode = None  # "q" or "a" — which buffer subsequent non-marker lines append to
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        q_match = re.match(r"^q\s*[:\-.]\s*(.*)", line, re.IGNORECASE)
        a_match = re.match(r"^a\s*[:\-.]\s*(.*)", line, re.IGNORECASE)
        if q_match:
            flush()
            current_q, current_a = q_match.group(1), None
            mode = "q"
        elif a_match:
            current_a = a_match.group(1)
            mode = "a"
        elif mode == "q" and current_q is not None:
            current_q += " " + line
        elif mode == "a" and current_a is not None:
            current_a += " " + line
    flush()
    return pairs


@app.post("/eval/parse-qa-pdf")
def eval_parse_qa_pdf(file: UploadFile | None = File(default=None)):
    """
    Accepts an uploaded PDF, TXT, or MD file containing Q:/A: formatted
    pairs, extracts its text (via PyMuPDF for PDF, direct decode for
    plain text), and returns the parsed pairs as JSON for the eval page
    to populate its question rows from — so a pre-written test set
    doesn't need to be retyped by hand into the UI.
    """
    if not file or not file.filename:
        return JSONResponse({"error": "No file uploaded"}, status_code=400)

    filename_lower = file.filename.lower()
    if filename_lower.endswith(".pdf"):
        tmp_path = Path(tempfile.gettempdir()) / f"qa_upload_{uuid.uuid4().hex}.pdf"
        try:
            _save_upload_to(file, tmp_path)
            pages = extract_pdf_pages(str(tmp_path))
            if not pages:
                return JSONResponse({"error": "Could not extract any text from that PDF — "
                                              "it may be corrupt, encrypted, or a scanned "
                                              "image without a text layer"}, status_code=400)
            full_text = "\n".join(p["text"] for p in pages)
        finally:
            tmp_path.unlink(missing_ok=True)
    elif filename_lower.endswith((".txt", ".md")):
        try:
            full_text = file.file.read().decode("utf-8", errors="replace")
        except Exception:
            return JSONResponse({"error": "Could not read that file as text"}, status_code=400)
    else:
        return JSONResponse({"error": "Only PDF, TXT, or MD files are supported here"}, status_code=400)

    pairs = parse_qa_pairs(full_text)
    if not pairs:
        return JSONResponse({"error": 'No "Q:"/"A:" pairs found. Expected format: '
                                      '"Q: your question" on one line, "A: expected '
                                      'answer" on the next, blank line between pairs.'},
                            status_code=400)
    return {"ok": True, "pairs": pairs}


@app.get("/eval", include_in_schema=False)
def eval_page(request: Request, sid: SessionId):
    """Serves the evaluation page via the React Single Page Application."""
    index_html = FRONTEND_DIST / "index.html"
    if index_html.exists():
        return FileResponse(str(index_html))
    return JSONResponse(
        {"message": "React frontend not built. Run 'npm run build' in the frontend/ directory."},
        status_code=503,
    )


# Named presets for the Week 4 ablation ladder — each maps to a specific,
# real combination of techniques already implemented, so "did adding X
# actually help" can be measured directly rather than inferred.
EVAL_PRESETS = {
    "tfidf":                {"force_tfidf": True,  "mode": None,            "rerank": False, "rewrite": False},
    "bm25-qdrant-blend":    {"force_tfidf": False, "mode": "hybrid-legacy", "rerank": False, "rewrite": False},
    "bm25-qdrant-rrf":      {"force_tfidf": False, "mode": "hybrid",        "rerank": False, "rewrite": False},
    "rrf-rerank":           {"force_tfidf": False, "mode": "hybrid",        "rerank": True,  "rewrite": False},
    "rrf-rerank-rewrite":   {"force_tfidf": False, "mode": "hybrid",        "rerank": True,  "rewrite": True},
}


def _retrieve_for_eval(active_store, query: str, k: int, mode: str, force_tfidf: bool = False) -> list[dict]:
    """
    Runs retrieval through the exact same functions /ask uses for each
    mode, so evaluation measures the REAL system's behavior rather than
    a separate reimplementation that could silently drift from it.

    force_tfidf deliberately bypasses embeddings even if configured —
    needed for the "tfidf" baseline preset to be a real, comparable
    baseline against a document that DOES have embeddings, rather than
    only ever being available when no API key exists at all.
    """
    has_embeddings = (_embeddings_configured() and active_store.vectors
                      and len(active_store.vectors) == len(active_store.chunks))
    if has_embeddings and not force_tfidf:
        if mode == "hybrid-legacy":
            return hybrid_search(active_store, query, top_k=k)
        if mode == "embed":
            return active_store.query(embed_text(query), top_k=k, min_score=0.0)
        return reciprocal_rank_fusion(active_store, query, top_k=k)  # "hybrid" (RRF)
    index = active_store.get_tfidf_index()
    return search_chunks(query, active_store.chunks, index, top_k=k)


def _hit_check(retrieved: list[dict], expected: str) -> bool:
    """A 'hit' means the expected section/filename text appears as a
    substring of any retrieved chunk's section or filename — matches the
    same matching style used throughout this project's manual testing
    (e.g. "5.2 Policy on Annual Leave" matching that exact section)."""
    if not expected:
        return False
    expected_lower = expected.lower()
    for r in retrieved:
        section = (r.get("section") or "").lower()
        filename = (r.get("filename") or "").lower()
        if expected_lower in section or expected_lower in filename:
            return True
    return False


def _rr_rank(retrieved: list[dict], expected: str = "", expected_doc: str = "", expected_section: str = "") -> tuple[bool, float, int]:
    exp_lower = (expected or "").strip().lower()
    exp_doc_lower = (expected_doc or "").strip().lower()
    exp_sec_lower = (expected_section or "").strip().lower()

    if not (exp_lower or exp_doc_lower or exp_sec_lower):
        return False, 0.0, 0

    for idx, r in enumerate(retrieved, start=1):
        section = (r.get("section") or "").lower()
        filename = (r.get("filename") or "").lower()
        text = (r.get("text") or "").lower()

        hit = False
        if exp_doc_lower and exp_doc_lower in filename:
            hit = True
        elif exp_sec_lower and (exp_sec_lower in section or exp_sec_lower in text):
            hit = True
        elif exp_lower and (exp_lower in section or exp_lower in filename or exp_lower in text):
            hit = True

        if hit:
            return True, 1.0 / idx, idx

    return False, 0.0, 0


def _run_eval_preset(active_store, q_id: str, question: str, expected: str, k: int, preset: dict,
                     expected_doc: str = "", expected_section: str = "") -> dict:
    """
    Runs ONE question through ONE named preset's full pipeline —
    optional query rewriting, retrieval, optional reranking — and
    returns a per-stage breakdown with question ID and ground-truth validation.
    """
    search_q = rewrite_query(question) if preset.get("rewrite") else question
    retrieved = _retrieve_for_eval(active_store, search_q, k, preset.get("mode"), preset.get("force_tfidf", False))
    rerank_score = None
    if preset.get("rerank") and len(retrieved) > 1:
        retrieved, rerank_score = rerank_with_llm(question, retrieved)
    hit, rr, rank = _rr_rank(retrieved, expected=expected, expected_doc=expected_doc, expected_section=expected_section)
    failure_type = "Success" if hit else "Retrieval Failure"
    return {
        "id": q_id,
        "question": question,
        "search_query": search_q if search_q != question else None,
        "expected": expected or expected_section or expected_doc,
        "expected_doc": expected_doc or None,
        "expected_section": expected_section or None,
        "hit": hit,
        "reciprocal_rank": round(rr, 4),
        "rank": rank,
        "failure_type": failure_type,
        "rerank_score": rerank_score,
        "retrieved": [
            {
                "section": r.get("section"), "filename": r.get("filename"),
                "score": round(r.get("score", 0.0), 4),
                "embed_score": round(r["embed_score"], 4) if "embed_score" in r else None,
                "bm25_score": round(r["keyword_score"], 4) if "keyword_score" in r else None,
            }
            for r in retrieved
        ],
    }


@app.post("/eval/run")
def eval_run(sid: RequiredSessionId, payload: EvalRunRequest | None = Body(default=None)):
    """
    Runs a real hit-rate@k evaluation: for each (question, expected) pair,
    retrieves top-k through the real retrieval pipeline and checks ground truth.
    """
    payload = payload or EvalRunRequest()
    k = max(1, min(payload.top_k or payload.k, 20))
    chunk_mode_filter = (payload.strategy_filter or payload.chunk_mode or "").strip() or None
    preset_names = payload.presets or list(EVAL_PRESETS.keys())
    legacy_modes = payload.modes

    questions = []
    for idx, q in enumerate(payload.questions):
        q_text = (q.question or "").strip()
        expected = (q.expected or "").strip()
        expected_doc = (q.expected_doc or "").strip()
        expected_section = (q.expected_section or "").strip()
        q_id = str(q.id or f"q_{idx + 1}")

        if q_text and (expected or expected_doc or expected_section):
            questions.append({
                "id": q_id,
                "question": q_text,
                "expected": expected or expected_section or expected_doc,
                "expected_doc": expected_doc,
                "expected_section": expected_section,
            })

    if not questions:
        return JSONResponse(
            {"error": "No valid questions provided (both question and expected ground truth are required)"},
            status_code=400,
        )

    store = _get_store(sid)
    if not store.chunks:
        return JSONResponse({"error": "No documents indexed in this session yet — upload one first"},
                            status_code=400)

    active_store = store.filtered_by_method(chunk_mode_filter) if chunk_mode_filter else store
    if chunk_mode_filter and not active_store.chunks:
        return JSONResponse({"error": f"No documents indexed under the '{chunk_mode_filter}' strategy"},
                            status_code=400)

    by_preset = {}
    names = legacy_modes if legacy_modes else preset_names
    for name in names:
        preset = EVAL_PRESETS.get(name) or {
            "force_tfidf": False, "mode": name, "rerank": False, "rewrite": False,
        }
        per_question = []
        hits = 0
        for q in questions:
            result = _run_eval_preset(
                active_store,
                q["id"],
                q["question"],
                q["expected"],
                k,
                preset,
                expected_doc=q["expected_doc"],
                expected_section=q["expected_section"]
            )
            hits += int(result["hit"])
            per_question.append(result)

        rr_sum = sum(r.get("reciprocal_rank", 0.0) for r in per_question)
        total = len(questions) or 1
        by_preset[name] = {
            "hit_rate": hits / total,
            "mrr": round(rr_sum / total, 4),
            "hits": hits,
            "total": total,
            "results": per_question,
        }

    return {"ok": True, "k": k, "total_questions": len(questions), "modes": by_preset}


@app.get("/healthz", response_model=HealthzResponse)
def healthz():
    """Liveness endpoint for deployment monitoring."""
    return HealthzResponse(
        status="ok",
        embeddings_configured=_embeddings_configured(),
        chat_configured=_chat_configured(),
        chat_backend=CHAT_BACKEND,
        embeddings_backend=EMBED_BACKEND,
        retrieval_mode=RETRIEVAL_MODE if _embeddings_configured() else "tfidf-only",
        vector_backend=VECTOR_BACKEND,
        active_sessions=len(SESSION_ACCESS),
    )


@app.get("/readyz", response_model=ReadyzResponse)
def readyz(response: Response):
    """Readiness endpoint: verifies external dependencies are reachable without expensive LLM inference."""
    checks = {"app": True}
    status_code = 200

    if VECTOR_BACKEND == "qdrant":
        try:
            from qdrant_store import _client
            _client().get_collections()
            checks["qdrant"] = True
        except Exception as exc:
            logger.warning("Readiness check: Qdrant unreachable: %s", exc)
            checks["qdrant"] = False
            status_code = 503

    response.status_code = status_code
    return ReadyzResponse(ready=status_code == 200, checks=checks)


@app.get("/status", response_model=StatusResponse)
def status(sid: OptionalSessionId):
    chunks = _get_store(sid).chunks if sid else []

    session_files = SESSION_FILES.get(sid, {})

    docs_seen: dict[str, dict] = {}
    for c in chunks:
        if c["doc_id"] not in docs_seen:
            docs_seen[c["doc_id"]] = {
                "filename": c["filename"],
                "doc_id": c["doc_id"],
                "chunk_count": 0,
                "method": c["method"],
                "openable": c["doc_id"] in session_files,
            }
        docs_seen[c["doc_id"]]["chunk_count"] += 1

    return StatusResponse(
        total_chunks=len(chunks),
        documents=list(docs_seen.values()),
        methods=sorted({c["method"] for c in chunks}),
        mode=(RETRIEVAL_MODE if _embeddings_configured() else "tfidf-only"),
        vector_backend=VECTOR_BACKEND,
    )


@app.get("/orphans")
def list_orphans(
    request: Request,
    sid: OptionalSessionId,
    session_id: str | None = Query(default=None,
                                   description="Admin-only: filter orphans to one session id."),
):
    """Operational reconciliation endpoint: returns currently unresolved orphaned documents.
    Enforces authorization:
      - Admin (via ADMIN_API_KEY in Bearer token or X-Admin-Key header): can view system-wide
        orphans or filter by ?session_id=..., with full operational metadata including stored_path.
      - Normal user (via active session): can view only their own session's orphans, with stored_path
        redacted for privacy.
      - Unauthenticated caller without valid session or admin credentials: rejected with 401 Unauthorized."""
    is_admin = _is_admin_request(request)

    if not is_admin and not sid:
        return JSONResponse({
            "error": "Unauthorized: active session or admin credentials required."
        }, status_code=401)

    durable_orphans = _read_durable_orphans()

    if is_admin:
        if session_id:
            raw_records = list(durable_orphans.get(session_id, []))
        else:
            raw_records = [item for items in durable_orphans.values() for item in items]
        return {
            "count": len(raw_records),
            "admin": True,
            "orphans": raw_records,
        }

    # Normal session: strictly isolated to caller's session, stored_path and internal error details redacted
    records = list(durable_orphans.get(sid, []))
    safe_records = [
        {
            "doc_id": r.get("doc_id"),
            "filename": r.get("filename"),
            "error": "Vector store cleanup failed",
            "timestamp": r.get("timestamp"),
            "status": r.get("status", "orphaned"),
        }
        for r in records
    ]
    return {
        "count": len(safe_records),
        "admin": False,
        "orphans": safe_records,
    }


# NOTE: the 413 "File too large" response is produced by MaxBodySizeMiddleware
# (see App Setup) rather than an exception handler — Starlette never raises a
# 413 of its own, because it imposes no body-size limit in the first place.


if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", 5000))
    # Defaults to loopback-only — safe for `python app.py` on your own
    # machine (not reachable from your LAN). MUST be "0.0.0.0" when run
    # inside Docker: a container's own loopback interface is invisible to
    # `docker run -p`/docker-compose's port mapping, so binding to
    # 127.0.0.1 in a container means the port mapping silently does
    # nothing — the container starts fine, but nothing outside it can
    # ever reach the server. docker-compose.yml sets HOST=0.0.0.0 for
    # exactly this reason; don't set it that way for a bare local run
    # unless you actually want this reachable from your whole network.
    host = os.environ.get("HOST", "127.0.0.1")
    debug = os.environ.get("APP_DEBUG", "").lower() in ("1", "true", "yes")
    embed_label = "Ollama (local)" if EMBED_BACKEND == "ollama" else "Gemini"
    chat_label = "Ollama (local)" if CHAT_BACKEND == "ollama" else "xAI Grok"
    mode_label = (f"embeddings ({embed_label}) + LLM ({chat_label})" if (_embeddings_configured() and _chat_configured())
                  else "TF-IDF (offline fallback)")
    print(f"\n🚀 Ask My Docs is running → http://localhost:{port}")
    print(f"   Mode: {mode_label}")
    print(f"   API docs: http://localhost:{port}/docs\n")
    # Ensure React 18 + Vite frontend is built before starting server
    ensure_frontend_built()

    # APP_DEBUG turns on uvicorn's autoreloader, which needs an import string
    # rather than the app object (it re-imports the module in the child process).
    uvicorn.run(
        "app:app" if debug else app,
        host=host,
        port=port,
        reload=debug,
        log_level="debug" if debug else "info",
    )